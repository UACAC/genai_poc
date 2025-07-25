# services/document_service.py
import io
import uuid
import base64
import requests
from typing import List, Optional, Union
from docx import Document
import fitz
import markdown
import os
import mimetypes
from bs4 import BeautifulSoup
from services.rag_service import RAGService
from services.agent_service import AgentService
from services.llm_service import LLMService
from services.database import SessionLocal, ComplianceAgent

class TemplateParser:
    @staticmethod
    def extract_headings_from_docx(path: str) -> List[str]:
        doc = Document(path)
        return [p.text for p in doc.paragraphs
                if p.style.name.startswith("Heading") and p.text.strip()]

    @staticmethod
    def extract_headings_from_pdf(path: str, size_threshold: float = 16.0) -> List[str]:
        doc = fitz.open(path)
        headings = []
        for page in doc:
            for block in page.get_text("dict")["blocks"]:
                if block["type"] != 0: continue
                for line in block["lines"]:
                    size = line["spans"][0]["size"]
                    text = "".join(s["text"] for s in line["spans"]).strip()
                    if size >= size_threshold and text:
                        headings.append(text)
        return headings

    @staticmethod
    def extract_headings_from_html(html: str) -> List[str]:
        soup = BeautifulSoup(html, "html.parser")
        headings = []
        for level in range(1,7):
            for tag in soup.find_all(f"h{level}"):
                text = tag.get_text(strip=True)
                if text:
                    headings.append(text)
        return headings

    @staticmethod
    def extract_headings_from_markdown(md: str) -> List[str]:
        lines = md.splitlines()
        return [l.lstrip("# ").strip() for l in lines
                if l.startswith("#")]
        
    @staticmethod
    def extract_headings_from_text(text: str) -> List[str]:
        return [
            line.lstrip("# ").strip()
            for line in text.splitlines()
            if line.startswith("# ")
        ]

    @classmethod
    def extract(cls, path_or_str: str, is_path: bool=True) -> List[str]:
        if not is_path:
            return cls.extract_headings_from_text(path_or_str)

        # ext = os.path.splitext(path_or_str if is_path else "")[1].lower()
        ext = os.path.splitext(path_or_str)[1].lower()
        
        if ext == ".docx":
            return cls.extract_headings_from_docx(path_or_str)
        if ext == ".docx":
            return cls.extract_headings_from_docx(path_or_str)
        if ext == ".pdf":
            return cls.extract_headings_from_pdf(path_or_str)
        if ext in {".html", ".htm"}:
            html = open(path_or_str).read() if is_path else path_or_str
            return cls.extract_headings_from_html(html)
        if ext == ".md" or (not is_path and "\n" in path_or_str):
            md = open(path_or_str).read() if is_path else path_or_str
            return cls.extract_headings_from_markdown(md)
        raise ValueError(f"Unsupported template format: {ext!r}")
class DocumentService:
    def __init__(
        self,
        rag_service: RAGService,
        agent_service: AgentService,
        llm_service: LLMService,
        chroma_url: str,
        agent_api_url: str,
    ):
        self.rag = rag_service
        self.agent = agent_service
        self.llm = llm_service
        self.chroma_url = chroma_url.rstrip("/")
        self.agent_api = agent_api_url.rstrip("/")

    def _fetch_templates(self, collection: str) -> List[str]:
        resp = requests.get(f"{self.chroma_url}/documents",
                            params={"collection_name": collection})
        resp.raise_for_status()
        return resp.json().get("documents", [])

    def _retrieve_context(self, tmpl: str, sources: List[str], top_k: int) -> str:
        pieces = []
        for coll in sources:
            docs, ok = self.rag.get_relevant_documents(tmpl, coll)
            if ok:
                pieces += docs[:top_k]
        return "\n\n".join(pieces)

    def _invoke_agent(
        self,
        prompt: str,
        agent_id: int,
        collection: str,
        ) -> str:
        """
        Use your LLMService.query_model to do a RAG-enabled completion.
        We assume that in your DB you can look up the ComplianceAgent
        to get its `model_name`
        """
        session = SessionLocal()
        try:
            agent = session.query(ComplianceAgent).get(agent_id)
            model_name = agent.model_name.lower()
        finally:
            session.close()

        # ask your LLMService to do RAG over `collection` + `prompt`
        answer, _ = self.llm.query_model(
            model_name=model_name,
            query=prompt,
            collection_name=collection,
            query_type="rag",
        )
        return answer
    
    def _reconstruct_template(self, doc_id: str, collection: str) -> str:
        resp = requests.get(
            f"{self.chroma_url}/documents/reconstruct/{doc_id}",
            params={"collection_name": collection}
        )
        resp.raise_for_status()

        data = resp.json()
        text = data.get("reconstructed_content")
        if text is None:
            raise RuntimeError(f"No reconstructed_content in response: {list(data.keys())!r}")

        return text


    def generate_documents(
        self,
        template_collection: str,
        template_doc_ids:    Optional[List[str]] = None,
        source_collections:  Optional[List[str]] = None,
        source_doc_ids:      Optional[List[str]] = None,
        agent_ids:           List[int]           = [],
        use_rag:             bool                = True,
        top_k:               int                 = 5,
        doc_title:          Optional[str]        = None,
    ) -> List[dict]:
        out = []
        # load agents once
        self.agent.load_selected_compliance_agents(agent_ids)

        # get raw markdown/text for each template
        tpl_texts = [
            self._reconstruct_template(tid, template_collection)
            for tid in (template_doc_ids or [])
        ]

        for tpl_text in tpl_texts:
            # extract all Markdown headings (#, ##, ###, etc.)
            headings = TemplateParser.extract_headings_from_markdown(tpl_text)

            # create a fresh .docx
            doc = Document()
            title = doc_title or "Generated Test Plan"
            doc.add_heading(title, level=1)

            # now fill in each section
            for heading in headings:
                # 1) RAG‐retrieve your source context for this heading
                ctx_pieces = []
                if use_rag and source_collections and source_doc_ids:
                    for coll, sid in zip(source_collections, source_doc_ids):
                        docs, ok = self.rag.get_relevant_documents(sid, coll)
                        if ok:
                            ctx_pieces += docs[:top_k]
                context = "\n\n".join(ctx_pieces)

                # 2) build the per-section prompt
                prompt = f"""### Section: {heading}

    Using the following source material, write the content for this section of the test plan:

    {context}
    """

                # 3) call your agent
                agent_meta = next(a for a in self.agent.compliance_agents if a["id"] == agent_ids[0])
                result = self.agent._invoke_chain(
                    agent=agent_meta,
                    data_sample=prompt,
                    session_id=str(uuid.uuid4()),
                    db=SessionLocal()
                )
                response_md = result["reason"]

                # 4) convert the Markdown response into *real* docx
                html = markdown.markdown(response_md)
                soup = BeautifulSoup(html, "html.parser")
                doc.add_heading(heading, level=2)
                for el in soup.contents:
                    if el.name and el.name.startswith("h"):
                        lvl = int(el.name[1])
                        doc.add_heading(el.get_text(), level=lvl)
                    elif el.name == "p":
                        doc.add_paragraph(el.get_text())
                    elif el.name == "ul":
                        for li in el.find_all("li"):
                            doc.add_paragraph(f"• {li.get_text()}")
                    elif el.name == "ol":
                        for idx, li in enumerate(el.find_all("li"), 1):
                            doc.add_paragraph(f"{idx}. {li.get_text()}")

            # 5) serialize to base64 & collect
            buf = io.BytesIO()
            doc.save(buf)
            out.append({
                "title":    title,
                "docx_b64": base64.b64encode(buf.getvalue()).decode("utf-8")
            })

        return out
