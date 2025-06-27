import os
import re
import tempfile
import time
import requests
from PyPDF2 import PdfReader
from sentence_transformers import SentenceTransformer
from docx import Document
from docx.shared import Inches
from fpdf import FPDF  
import streamlit as st
import base64
from io import BytesIO
from PIL import Image
import datetime
from io import BytesIO
from concurrent.futures import ThreadPoolExecutor, as_completed


# ChromaDB API endpoint
CHROMADB_API = os.getenv("CHROMA_URL", "http://localhost:8020")
LLM_API = os.getenv("LLM_API", "http://localhost:9020")

# Load Sentence Transformer model
embedding_model = SentenceTransformer('multi-qa-mpnet-base-dot-v1')

def fetch_collections():
    try:
        response = requests.get(f"{CHROMADB_API}/collections", timeout=10)
        response.raise_for_status()
        return response.json().get("collections", [])
    except requests.exceptions.RequestException as e:
        print(f"Fetch collections failed: {e}")
        return []

def get_available_models():
    try:
        response = requests.get(f"{LLM_API}/health", timeout=10)
        if response.status_code == 200:
            health_data = response.json()
            if "models" in health_data:
                return [model for model, status in health_data["models"].items() if "available" in status.lower()]
    except Exception as e:
        print(f"Error fetching models: {e}")
    return ["gpt-4", "gpt-3.5-turbo", "llama3"]

def check_model_availability():
    available_models = []
    test_models = ["gpt-4", "gpt-3.5-turbo", "llama3"]
    for model in test_models:
        try:
            response = requests.post(f"{LLM_API}/chat", json={"query": "Hello", "model": model, "use_rag": False}, timeout=30)
            if response.status_code == 200:
                available_models.append(model)
        except Exception as e:
            print(f"{model} error: {e}")
    return available_models

@st.cache_data(ttl=300)
def get_available_models_cached():
    return get_available_models()

def extract_sections_from_pdf(pdf_path):
    reader = PdfReader(pdf_path)
    full_text = "".join([page.extract_text() or "" for page in reader.pages])
    return [s.strip() for s in re.split(r'\n(?=\d+\.\s)', full_text) if s.strip()]

def extract_text_from_txt(txt_path):
    with open(txt_path, 'r', encoding="utf-8") as file:
        return [s.strip() for s in file.read().split("\n\n") if s.strip()]

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    return [s.strip() for s in re.split(r'\n(?=[A-Z])', "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])) if s.strip()]

def store_files_in_chromadb(files, collection_name, model_name="none", openai_api_key=None, 
                            chunk_size=1000, chunk_overlap=200, store_images=True, 
                            debug_mode=False, run_all_vision_models=True):
    """Store uploaded files in ChromaDB with enhanced parameters"""
    files_data = [('files', (file.name, file.getvalue(), file.type)) for file in files]
    
    params = {
        "collection_name": collection_name,
        "chunk_size": chunk_size,
        "chunk_overlap": chunk_overlap,
        "store_images": store_images,
        "model_name": model_name,
        "debug_mode": debug_mode,
        "run_all_vision_models": run_all_vision_models  # Add new parameter
    }
    
    # Prepare headers
    headers = {}
    if openai_api_key and model_name in ["gpt-4o-mini"]:
        headers["X-OpenAI-API-Key"] = openai_api_key
    
    try:
        response = requests.post(
            f"{CHROMADB_API}/documents/upload-and-process", 
            params=params, 
            files=files_data, 
            headers=headers, 
            timeout=600  # Increased timeout for multi-model processing
        )
        
        if response.status_code != 200:
            raise Exception(f"Failed to store documents: {response.status_code} - {response.text}")
        
        return response.json()
        
    except requests.exceptions.Timeout:
        raise Exception("Request timed out. Multi-model processing takes longer but provides comprehensive analysis.")
    except requests.exceptions.ConnectionError:
        raise Exception(f"Could not connect to ChromaDB at {CHROMADB_API}")
    except Exception as e:
        raise Exception(f"Error storing files: {str(e)}")
    
def store_files_in_chromadb_selective(files, collection_name, model_name="none", openai_api_key=None, 
                                    chunk_size=1000, chunk_overlap=200, store_images=True, 
                                    debug_mode=False, selected_models=None):
    """Store uploaded files in ChromaDB with selective vision models"""
    files_data = [('files', (file.name, file.getvalue(), file.type)) for file in files]
    
    # Default to enhanced + basic if no models selected
    if not selected_models:
        selected_models = ["enhanced_local", "basic"]
    
    # Convert list to comma-separated string
    vision_models_str = ",".join(selected_models)
    
    params = {
        "collection_name": collection_name,
        "chunk_size": chunk_size,
        "chunk_overlap": chunk_overlap,
        "store_images": store_images,
        "model_name": model_name,
        "debug_mode": debug_mode,
        "vision_models": vision_models_str  # Pass selected models
    }
    
    # Prepare headers
    headers = {}
    if openai_api_key and model_name in ["gpt-4o-mini"]:
        headers["X-OpenAI-API-Key"] = openai_api_key
    
    try:
        # Adjust timeout based on number of models selected
        base_timeout = 300
        model_timeout = len(selected_models) * 60  # 1 minute per model
        total_timeout = min(base_timeout + model_timeout, 900)  # Max 15 minutes
        
        response = requests.post(
            f"{CHROMADB_API}/documents/upload-and-process", 
            params=params, 
            files=files_data, 
            headers=headers, 
            timeout=total_timeout
        )
        
        if response.status_code != 200:
            raise Exception(f"Failed to store documents: {response.status_code} - {response.text}")
        
        return response.json()
        
    except requests.exceptions.Timeout:
        raise Exception(f"Request timed out. Processing {len(selected_models)} vision models takes significant time.")
    except requests.exceptions.ConnectionError:
        raise Exception(f"Could not connect to ChromaDB at {CHROMADB_API}")
    except Exception as e:
        raise Exception(f"Error storing files: {str(e)}")


def store_files_in_chromadb_parallel(
    files,
    collection_name,
    model_name="none",
    openai_api_key=None,
    chunk_size=1000,
    chunk_overlap=200,
    store_images=True,
    debug_mode=False,
    selected_models=None,
    max_workers=4
):
    """
    Fire off one HTTP upload per file in parallel threads, then gather all the responses.
    """
    # worker that wraps your existing single‐file call
    def _upload_one(f):
        return store_files_in_chromadb_selective(
            [f],
            collection_name,
            model_name=model_name,
            openai_api_key=openai_api_key,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            store_images=store_images,
            debug_mode=debug_mode,
            selected_models=selected_models,
        )

    results = []
    with ThreadPoolExecutor(max_workers=max_workers) as pool:
        futures = {pool.submit(_upload_one, f): f.name for f in files}
        for fut in as_completed(futures):
            fname = futures[fut]
            try:
                results.append(fut.result())
            except Exception as e:
                results.append({
                    "filename": fname,
                    "status": "error",
                    "error": str(e)
                })
                
    merged = []
    for resp in results:
        if isinstance(resp, dict) and "processed_files" in resp:
            merged.extend(resp["processed_files"])
        else:
            merged.append(resp)
    
    return {
        "collection": collection_name,
        "processed_files": merged,
        "total_files_processed": sum(1 for d in merged if d.get("status")=="success")
    }


def query_documents_with_embedding(collection_name, query_text, n_results=5):
    """Query documents using text embedding"""
    try:
        # Generate embedding for the query
        query_embedding = embedding_model.encode([query_text]).tolist()
        
        # Query ChromaDB
        response = requests.post(
            f"{CHROMADB_API}/documents/query",
            json={
                "collection_name": collection_name,
                "query_embeddings": query_embedding,
                "n_results": n_results,
                "include": ["documents", "metadatas", "distances"]
            },
            timeout=60
        )
        
        if response.status_code == 200:
            return response.json()
        else:
            raise Exception(f"Query failed: {response.text}")
            
    except Exception as e:
        raise Exception(f"Error querying documents: {str(e)}")

def get_all_documents_in_collection(collection_name):
    """Get all unique documents in a collection with their metadata"""
    try:
        response = requests.get(
            f"{CHROMADB_API}/documents",
            params={"collection_name": collection_name},
            timeout=60
        )
        
        if response.status_code != 200:
            raise Exception(f"Failed to fetch documents: {response.text}")
            
        data = response.json()
        
        # Group chunks by document_id to get unique documents
        documents = {}
        for i, doc_id in enumerate(data.get("ids", [])):
            metadata = data["metadatas"][i] if i < len(data.get("metadatas", [])) else {}
            document_id = metadata.get("document_id")
            document_name = metadata.get("document_name", "Unknown")
            
            if document_id and document_id not in documents:
                documents[document_id] = {
                    "document_id": document_id,
                    "document_name": document_name,
                    "file_type": metadata.get("file_type", ""),
                    "total_chunks": metadata.get("total_chunks", 0),
                    "has_images": metadata.get("has_images", False),
                    "image_count": metadata.get("image_count", 0),
                    "processing_timestamp": metadata.get("timestamp", ""),
                    "openai_api_used": metadata.get("openai_api_used", False)
                }
        
        return list(documents.values())
        
    except Exception as e:
        raise Exception(f"Error fetching documents: {str(e)}")

def list_all_chunks_with_scores(collection_name, query_text=None):
    """List all chunks with optional similarity scores"""
    response = requests.get(f"{CHROMADB_API}/documents", params={"collection_name": collection_name})
    if response.status_code != 200:
        return []
    
    docs = response.json()
    scores_dict = {}
    
    if query_text:
        query_embedding = embedding_model.encode([query_text]).tolist()
        score_response = requests.post(
            f"{CHROMADB_API}/documents/query", 
            json={
                "collection_name": collection_name, 
                "query_embeddings": query_embedding, 
                "n_results": len(docs["ids"]), 
                "include": ["metadatas", "distances"]
            }
        )
        if score_response.status_code == 200:
            results = score_response.json()
            scores_dict = {doc_id: round(distance, 4) for doc_id, distance in zip(results["ids"][0], results["distances"][0])}
    
    return [
        {
            "Collection": collection_name, 
            "Document Name": metadata.get("document_name", "Unknown"), 
            "Document ID": doc_id, 
            "Chunk ID": f"`{doc_id}`", 
            "Document Text": f">{doc_text[:250] + '...' if len(doc_text) > 250 else doc_text}", 
            "Metadata": f"**{(metadata or {}).get('document_name', 'Unknown')}**", 
            "Score": scores_dict.get(doc_id, "N/A"),
            "Has Images": metadata.get("has_images", False),
            "Image Count": metadata.get("image_count", 0)
        } 
        for doc_id, doc_text, metadata in zip(
            docs["ids"], 
            docs["documents"], 
            docs.get("metadatas", [{}] * len(docs["ids"]))
        ) 
    ]

def reconstruct_document_with_timeout(document_id, collection_name, timeout=300):
    """Reconstruct document with configurable timeout"""
    try:
        response = requests.get(
            f"{CHROMADB_API}/documents/reconstruct/{document_id}",
            params={"collection_name": collection_name},
            timeout=timeout
        )
        
        if response.status_code == 200:
            return response.json()
        elif response.status_code == 404:
            raise Exception("Document not found")
        else:
            raise Exception(f"Error: {response.text}")
            
    except requests.exceptions.Timeout:
        raise Exception("Request timed out. The document might be very large or the server is busy.")
    except Exception as e:
        raise Exception(f"Error reconstructing document: {str(e)}")

def image_to_base64(img_url):
    """Convert image URL to base64 for embedding in markdown"""
    try:
        response = requests.get(img_url, timeout=30)
        if response.status_code == 200:
            img = Image.open(BytesIO(response.content))
            buffered = BytesIO()
            img.save(buffered, format="PNG")
            img_base64 = base64.b64encode(buffered.getvalue()).decode()
            return f"![img](data:image/png;base64,{img_base64})"
        else:
            return f"Image not found: {img_url}"
    except Exception as e:
        return f"Failed to load image: {img_url} — {str(e)}"

def get_image_from_chromadb(filename):
    """Get image from ChromaDB storage"""
    try:
        img_url = f"{CHROMADB_API}/images/{filename}"
        response = requests.get(img_url, timeout=30)
        
        if response.status_code == 200:
            return Image.open(BytesIO(response.content))
        else:
            return None
    except Exception as e:
        print(f"Error fetching image {filename}: {e}")
        return None

def render_reconstructed_document(result: dict):
    """Render reconstructed document with enhanced image handling"""
    rich_content = result.get("reconstructed_content", "")
    images = result.get("images", [])

    # Replace image markers with base64 images + description
    for image in images:
        marker = f"[IMAGE:{image['filename']}]"
        img_url = f"{CHROMADB_API}/images/{image['filename']}"
        base64_img = image_to_base64(img_url)
        description = image.get("description", "")
        
        # Enhanced replacement with better formatting
        replacement = f"""
{base64_img}

**Image Analysis:** {description}

---
"""
        rich_content = rich_content.replace(marker, replacement)

    # Display the content
    st.markdown(rich_content, unsafe_allow_html=True)

def clean_text(text):
    """Remove null bytes and non-XML-safe characters"""
    if not text:
        return ""
    return ''.join(c for c in text if c == '\n' or c == '\t' or (32 <= ord(c) <= 126))

def export_to_docx(result):
    """Export document to DOCX with enhanced image handling"""
    return export_to_docx_with_markdown(result)

def export_to_pdf(result):
    """Export document to PDF with enhanced image handling"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    content = result["reconstructed_content"]
    images = result.get("images", [])
    image_map = {img["filename"]: img["description"] for img in images}

    # Split content by image markers
    parts = content.split("[IMAGE:")
    
    # Add first part (before any images)
    if parts[0].strip():
        pdf.multi_cell(0, 10, clean_text(parts[0].strip()))

    # Process each part that contains an image
    for part in parts[1:]:
        if "]" in part:
            filename, rest = part.split("]", 1)
            
            # Try to get and insert the image
            image_obj = get_image_from_chromadb(filename)
            if image_obj:
                try:
                    # Save image to temporary file
                    temp_img = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                    image_obj.save(temp_img.name)

                    # Add new page for image
                    pdf.add_page()
                    pdf.image(temp_img.name, w=150)
                    pdf.ln(5)
                    
                    # Add image description
                    if filename in image_map:
                        pdf.multi_cell(0, 10, clean_text(image_map[filename]))
                    
                    # Clean up temp file
                    os.unlink(temp_img.name)
                    
                except Exception as e:
                    pdf.multi_cell(0, 10, f"[Failed to insert image {filename}]: {e}")
            else:
                pdf.multi_cell(0, 10, f"[Image {filename} not found]")

            # Add remaining text after image
            if rest.strip():
                pdf.multi_cell(0, 10, clean_text(rest.strip()))
        else:
            # No image marker found, just add text
            if part.strip():
                pdf.multi_cell(0, 10, clean_text(part.strip()))

    # Save to temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    pdf.output(temp_file.name)
    return temp_file.name

def check_chromadb_health():
    """Check ChromaDB health and return status info"""
    try:
        response = requests.get(f"{CHROMADB_API}/health", timeout=10)
        if response.status_code == 200:
            return True, response.json()
        return False, None
    except Exception as e:
        return False, str(e)
    
# Add these additional functions to your utils.py file

def create_file_like_object(content, filename):
    """Create a file-like object from text content for upload"""

    
    # Encode content as bytes
    content_bytes = content.encode('utf-8')
    
    # Create BytesIO object
    file_obj = BytesIO(content_bytes)
    file_obj.name = filename
    
    return file_obj

def store_generated_document(content, filename, collection_name, metadata=None):
    """Store a generated document directly in ChromaDB"""
    try:
        # Create file-like object
        file_obj = create_file_like_object(content, filename)
        
        # Store using existing function
        result = store_files_in_chromadb(
            [file_obj], 
            collection_name,
            model_name="basic",
            chunk_size=2000,
            chunk_overlap=200,
            store_images=False
        )
        
        return result
        
    except Exception as e:
        raise Exception(f"Error storing generated document: {str(e)}")

def batch_store_generated_documents(documents, collection_name):
    """Store multiple generated documents in ChromaDB
    
    Args:
        documents: List of dicts with 'content', 'filename', and optional 'metadata'
        collection_name: Target collection name
    
    Returns:
        Dict with success/failure counts and details
    """
    results = {
        'successful': 0,
        'failed': 0,
        'errors': []
    }
    
    for doc in documents:
        try:
            store_generated_document(
                doc['content'], 
                doc['filename'], 
                collection_name, 
                doc.get('metadata')
            )
            results['successful'] += 1
            
        except Exception as e:
            results['failed'] += 1
            results['errors'].append({
                'filename': doc['filename'],
                'error': str(e)
            })
    
    return results

def get_collection_statistics(collection_name):
    """Get detailed statistics for a collection"""
    try:
        documents = get_all_documents_in_collection(collection_name)
        
        if not documents:
            return {
                'total_documents': 0,
                'total_chunks': 0,
                'file_types': {},
                'has_images': 0,
                'avg_chunks_per_doc': 0
            }
        
        # Calculate statistics
        total_docs = len(documents)
        total_chunks = sum(doc.get('total_chunks', 0) for doc in documents)
        has_images = sum(1 for doc in documents if doc.get('has_images', False))
        
        # File type breakdown
        file_types = {}
        for doc in documents:
            file_type = doc.get('file_type', 'unknown')
            file_types[file_type] = file_types.get(file_type, 0) + 1
        
        return {
            'total_documents': total_docs,
            'total_chunks': total_chunks,
            'file_types': file_types,
            'has_images': has_images,
            'avg_chunks_per_doc': round(total_chunks / total_docs, 1) if total_docs > 0 else 0,
            'documents': documents
        }
        
    except Exception as e:
        raise Exception(f"Error getting collection statistics: {str(e)}")

def search_generated_documents(collection_name, search_term, max_results=10):
    """Search for generated documents containing specific terms"""
    try:
        # Use existing query function
        results = query_documents_with_embedding(collection_name, search_term, max_results)
        
        # Process results for generated documents
        if results and 'documents' in results:
            processed_results = []
            
            for i, (doc_id, content, metadata, distance) in enumerate(zip(
                results.get('ids', []),
                results.get('documents', []),
                results.get('metadatas', []),
                results.get('distances', [])
            )):
                processed_results.append({
                    'chunk_id': doc_id,
                    'content_preview': content[:300] + "..." if len(content) > 300 else content,
                    'document_name': metadata.get('document_name', 'Unknown'),
                    'document_id': metadata.get('document_id', 'Unknown'),
                    'similarity_score': round(1 - distance, 3),  # Convert distance to similarity
                    'file_type': metadata.get('file_type', 'unknown')
                })
            
            return processed_results
        
        return []
        
    except Exception as e:
        raise Exception(f"Error searching documents: {str(e)}")

def export_collection_summary(collection_name):
    """Export a summary of all documents in a collection"""
    try:
        stats = get_collection_statistics(collection_name)
        
        if stats['total_documents'] == 0:
            return "No documents found in collection"
        
        # Create summary text
        summary = f"""Collection Summary: {collection_name}
Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

Overview:
- Total Documents: {stats['total_documents']}
- Total Text Chunks: {stats['total_chunks']}
- Documents with Images: {stats['has_images']}
- Average Chunks per Document: {stats['avg_chunks_per_doc']}

File Type Breakdown:
"""
        
        for file_type, count in stats['file_types'].items():
            summary += f"- {file_type.upper()}: {count} documents\n"
        
        summary += "\nDocument List:\n"
        
        for i, doc in enumerate(stats['documents'], 1):
            summary += f"{i}. {doc['document_name']}\n"
            summary += f"   Type: {doc['file_type']} | Chunks: {doc['total_chunks']} | "
            summary += f"Images: {'Yes' if doc['has_images'] else 'No'}\n"
            summary += f"   ID: {doc['document_id']}\n"
            summary += f"   Uploaded: {doc['processing_timestamp'][:10] if doc['processing_timestamp'] else 'Unknown'}\n\n"
        
        return summary
        
    except Exception as e:
        return f"Error generating collection summary: {str(e)}"

def cleanup_temp_files(file_paths):
    """Clean up temporary files safely"""
    cleaned = 0
    errors = []
    
    for file_path in file_paths:
        try:
            if os.path.exists(file_path):
                os.unlink(file_path)
                cleaned += 1
        except Exception as e:
            errors.append(f"Failed to cleanup {file_path}: {e}")
    
    return cleaned, errors

def validate_generated_document(result_dict):
    """Validate that a generated document result has required fields"""
    required_fields = [
        'document_title', 'source_document', 'agent_name', 
        'analysis_content', 'generation_timestamp'
    ]
    
    missing_fields = []
    for field in required_fields:
        if field not in result_dict or not result_dict[field]:
            missing_fields.append(field)
    
    if missing_fields:
        raise ValueError(f"Generated document missing required fields: {missing_fields}")
    
    # Validate content length
    if len(result_dict['analysis_content']) < 10:
        raise ValueError("Generated content is too short")
    
    return True

def format_generation_metadata(result):
    """Format generation metadata for display"""
    return {
        "Document Title": result.get('document_title', 'Unknown'),
        "Source Document": result.get('source_document', 'Unknown'),
        "Generated by Agent": result.get('agent_name', 'Unknown'),
        "Agent ID": result.get('agent_id', 'Unknown'),
        "Generation Time": result.get('generation_timestamp', 'Unknown')[:19],
        "Content Length": f"{result.get('content_length', 0):,} characters",
        "Processing Time": f"{result.get('processing_time_ms', 0)}ms"
    }

def markdown_to_docx(markdown_text, doc):
    """Convert markdown text to properly formatted DOCX content"""
    lines = markdown_text.split('\n')
    for line in lines:
        l = line.strip()
        if not l:
            # Add empty paragraph for spacing
            doc.add_paragraph("")
            continue
            
        # Handle headers with ## 
        if l.startswith("## "):
            doc.add_heading(l.replace("##", "").strip(), level=1)
        elif l.startswith("### "):
            doc.add_heading(l.replace("###", "").strip(), level=2)
        elif l.startswith("#### "):
            doc.add_heading(l.replace("####", "").strip(), level=3)
        # Handle lines that are fully bold as headings
        elif l.startswith("**") and l.endswith("**") and l.count("**") == 2:
            doc.add_heading(l.replace("*", "").strip(), level=2)
        # Handle bullet points
        elif l.startswith(("-", "*", "•")):
            doc.add_paragraph(l.lstrip("-*• ").strip(), style='List Bullet')
        # Handle numbered lists
        elif l[:2].isdigit() and l[2:4] in ('.', ') '):
            doc.add_paragraph(l.split('.', 1)[1].strip() if '.' in l else l.split(')', 1)[1].strip(), style='List Number')
        # Handle lines with mixed bold text
        elif "**" in l:
            parts = l.split("**")
            p = doc.add_paragraph()
            toggle = False
            for part in parts:
                if part:  # Only add non-empty parts
                    run = p.add_run(part)
                    if toggle:
                        run.bold = True
                toggle = not toggle
        # Handle horizontal rules
        elif l.startswith("---"):
            doc.add_paragraph("_" * 50)  # Add a line separator
        # Regular paragraph
        else:
            doc.add_paragraph(l)

def export_to_docx_with_markdown(result):
    """Enhanced export to DOCX with markdown formatting support"""
    try:
        from docx import Document
        import tempfile
        
        doc = Document()
        doc.add_heading(result["document_name"], 0)

        rich_content = result["reconstructed_content"]
        images = result.get("images", [])

        # Create image lookup for descriptions
        image_lookup = {img['filename']: img for img in images}

        # Replace image markers with placeholders for processing
        for img in images:
            marker = f"[IMAGE:{img['filename']}]"
            if marker in rich_content:
                rich_content = rich_content.replace(marker, f"@@IMAGE:{img['filename']}@@")

        # Split content by image markers
        chunks = re.split(r'(@@IMAGE:.+?@@)', rich_content)

        for chunk in chunks:
            img_match = re.match(r'@@IMAGE:(.+?)@@', chunk)
            if img_match:
                filename = img_match.group(1)
                
                # Try to get and insert the image
                image_obj = get_image_from_chromadb(filename)
                if image_obj:
                    try:
                        # Save image to temporary file for docx
                        temp_img = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                        image_obj.save(temp_img.name)
                        
                        # Add image to document
                        from docx.shared import Inches
                        doc.add_picture(temp_img.name, width=Inches(5.5))
                        
                        # Add image description if available
                        if filename in image_lookup:
                            desc = image_lookup[filename].get('description', '')
                            if desc:
                                doc.add_paragraph(clean_text(desc), style='Intense Quote')
                        
                        # Clean up temp file
                        os.unlink(temp_img.name)
                        
                    except Exception as e:
                        doc.add_paragraph(f"[Image '{filename}' could not be inserted: {e}]")
                else:
                    doc.add_paragraph(f"[Image '{filename}' not found]")
            else:
                # Process text chunk with markdown formatting
                cleaned_chunk = clean_text(chunk.strip())
                if cleaned_chunk:
                    # KEY FIX: Use the markdown_to_docx function for proper formatting
                    markdown_to_docx(cleaned_chunk, doc)

        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)
        return temp_file.name
        
    except Exception as e:
        raise Exception(f"Error creating DOCX: {str(e)}")

def create_rule_document_docx(document_title, generated_content, metadata=None):
    """Create a DOCX file from generated rule content with markdown formatting"""
    try:
        # Create document
        doc = Document()
        doc.add_heading(document_title, 0)
        
        # Add metadata if provided
        if metadata:
            doc.add_heading('Document Information', 1)
            for key, value in metadata.items():
                p = doc.add_paragraph()
                p.add_run(f"{key}: ").bold = True
                p.add_run(str(value))
            
            # Add separator
            doc.add_paragraph("")
        
        # Add main content heading
        doc.add_heading('Analysis Content', 1)
        
        # Process the generated content with markdown formatting
        # This is the key fix - actually convert the markdown
        markdown_to_docx(generated_content, doc)
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)
        return temp_file.name
        
    except Exception as e:
        raise Exception(f"Error creating DOCX: {str(e)}")

def create_combined_rule_documents_docx(generated_documents, collection_name):
    """Create a combined DOCX file from multiple generated rule documents"""
    try:
        # Create combined document
        doc = Document()
        doc.add_heading('Generated Rule Documents Collection', 0)
        
        # Add summary information
        doc.add_heading('Collection Summary', 1)
        doc.add_paragraph(f"Collection Name: {collection_name}")
        doc.add_paragraph(f"Total Documents: {len(generated_documents)}")
        doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Add table of contents
        doc.add_heading('Table of Contents', 1)
        for i, doc_info in enumerate(generated_documents, 1):
            doc.add_paragraph(f"{i}. {doc_info['document_title']}")
        
        doc.add_page_break()
        
        # Add each document
        for i, doc_info in enumerate(generated_documents, 1):
            # Document header
            doc.add_heading(f"{i}. {doc_info['document_title']}", 1)
            
            # Document metadata
            metadata_info = [
                f"Source Document: {doc_info.get('source_document', 'Unknown')}",
                f"Generated by Agent: {doc_info.get('agent_name', 'Unknown')}",
                f"Rules Count: {doc_info.get('rules_count', 'Unknown')}",
                f"Content Length: {doc_info.get('content_length', 0):,} characters",
                f"Generated: {doc_info.get('generation_timestamp', 'Unknown')[:19]}"
            ]
            
            for info in metadata_info:
                doc.add_paragraph(info)
            
            doc.add_paragraph("")  # Empty line
            
            # Process document content with markdown formatting
            content = doc_info.get('generated_content', '') or doc_info.get('analysis_content', '')
            if content:
                # KEY FIX: Apply markdown conversion here
                markdown_to_docx(content, doc)
            
            # Add page break except for last document
            if i < len(generated_documents):
                doc.add_page_break()
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)
        return temp_file.name
        
    except Exception as e:
        raise Exception(f"Error creating combined DOCX: {str(e)}")

def export_generated_documents_as_docx(generated_documents, output_collection):
    """Export generated documents using the markdown to DOCX conversion"""
    try:
        # Create combined document
        docx_path = create_combined_rule_documents_docx(generated_documents, output_collection)
        
        # Read the file
        with open(docx_path, "rb") as file:
            file_data = file.read()
        
        # Clean up temp file
        os.unlink(docx_path)
        
        return file_data, f"Generated_Rules_{output_collection}_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        
    except Exception as e:
        raise Exception(f"Error exporting documents: {str(e)}")

def export_single_document_as_docx(document_info):
    """Export a single generated document as DOCX with markdown formatting"""
    try:
        # Prepare metadata
        metadata = {
            "Source Document": document_info.get('source_document', 'Unknown'),
            "Generated by Agent": document_info.get('agent_name', 'Unknown'),
            "Agent ID": document_info.get('agent_id', 'Unknown'),
            "Rules Count": document_info.get('rules_count', 'Unknown'),
            "Content Length": f"{document_info.get('content_length', 0):,} characters"
        }
        
        # Create DOCX
        docx_path = create_rule_document_docx(
            document_info['document_title'],
            document_info.get('generated_content', '') or document_info.get('analysis_content', ''),
            metadata
        )
        
        # Read the file
        with open(docx_path, "rb") as file:
            file_data = file.read()
        
        # Clean up temp file
        os.unlink(docx_path)
        
        # Create filename
        safe_title = "".join(c for c in document_info['document_title'] if c.isalnum() or c in (' ', '-', '_')).strip()
        filename = f"{safe_title}.docx"
        
        return file_data, filename
        
    except Exception as e:
        raise Exception(f"Error exporting single document: {str(e)}")
    
def build_docx_bytes(generated_results):
    # 1) Create a new Document
    doc = Document()
    doc.add_heading('Generated Analysis Documents', level=0)

    # 2) Iterate and add each result
    for idx, result in enumerate(generated_results, 1):
        doc.add_heading(f"{idx}. {result['document_title']}", level=1)
        # insert metadata if you like
        doc.add_paragraph(f"Source: {result['source_document']}")
        doc.add_paragraph(f"Agent: {result['agent_name']}")
        doc.add_paragraph('')  # blank line
        # assume markdown_to_docx is already imported and converts your markdown
        markdown_to_docx(result['analysis_content'], doc)
        if idx < len(generated_results):
            doc.add_page_break()

    # 3) Save into a BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer