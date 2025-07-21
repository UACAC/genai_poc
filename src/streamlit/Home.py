import streamlit as st
import requests
import os
import nest_asyncio
import datetime
import time
from utils import * 
import torch
import base64
from upload_documents import render_upload_component

torch.classes.__path__ = []
nest_asyncio.apply()

# THIS MUST BE THE VERY FIRST STREAMLIT COMMAND
st.set_page_config(page_title="AI Assistant", layout="wide")

# FastAPI API endpoints
FASTAPI_API = os.getenv("FASTAPI_URL", "http://localhost:9020")
CHROMADB_API = os.getenv("CHROMA_URL", "http://localhost:8020") 
CHAT_ENDPOINT = f"{FASTAPI_API}/chat"
HISTORY_ENDPOINT = f"{FASTAPI_API}/chat-history"
HEALTH_ENDPOINT = f"{FASTAPI_API}/health"
OPEN_AI_API_KEY = os.getenv("OPEN_AI_API_KEY")

st.title("AI Assistant")

# Initialize session state
if 'health_status' not in st.session_state:
    st.session_state.health_status = None
if 'available_models' not in st.session_state:
    st.session_state.available_models = []
if 'collections' not in st.session_state:
    st.session_state.collections = []
if 'agents_data' not in st.session_state:
    st.session_state.agents_data = []
if 'debate_sequence' not in st.session_state:
    st.session_state.debate_sequence = []
if 'upload_progress' not in st.session_state:
    st.session_state.upload_progress = {}

# Cache functions
@st.cache_data(ttl=1_200) #20 Minutes
def get_available_models_cached():
    return get_available_models()

def check_model_status(model_name):
    """Check if a specific model is loaded in Ollama"""
    try:
        response = requests.get(f"{FASTAPI_API}/health", timeout=5)
        if response.ok:
            health_data = response.json()
            models = health_data.get("models", {})
            return models.get(model_name, "unknown")
    except:
        return "unknown"

def upload_documents_to_chromadb(files, collection_name, openai_api_key=OPEN_AI_API_KEY):
    """Upload documents to ChromaDB using the existing endpoint"""
    try:
        # Prepare files for upload
        files_data = []
        for file in files:
            files_data.append(("files", (file.name, file.getvalue(), file.type)))
        
        # Prepare headers
        headers = {}
        if openai_api_key:
            headers["X-OpenAI-API-Key"] = openai_api_key
        
        # Prepare parameters
        params = {
            "collection_name": collection_name,
            "chunk_size": 1000,
            "chunk_overlap": 200,
            "store_images": True,
            "model_name": "enhanced",
            "debug_mode": False,
            "run_all_vision_models": True
        }
        
        # Make request to ChromaDB upload endpoint
        response = requests.post(
            f"{CHROMADB_API}/documents/upload-and-process",
            files=files_data,
            params=params,
            headers=headers,
            timeout=300
        )
        
        return response
        
    except Exception as e:
        raise Exception(f"Upload failed: {str(e)}")

# def create_collection(collection_name):
#     """Create a new ChromaDB collection"""
#     try:
#         response = requests.post(
#             f"{CHROMADB_URL}/collection/create",
#             params={"collection_name": collection_name},
#             timeout=30
#         )
#         return response
#     except Exception as e:
#         raise Exception(f"Collection creation failed: {str(e)}")

# def get_chromadb_collections():
#     """Get list of collections from ChromaDB"""
#     try:
#         response = requests.get(f"{CHROMADB_URL}/collections", timeout=10)
#         if response.ok:
#             return response.json().get("collections", [])
#         return []
#     except:
#         return []

# Model configurations
model_key_map = {
    "GPT-4": "gpt-4",
    "GPT-4": "gpt-3.5-turbo", 
    "LLaMA 3": "llama3",
}

model_descriptions = {
    "GPT-4": "Most capable model for complex analysis",
    "GPT-4": "Cost-effective model for general tasks",
    "LLaMA 3": "Fast and efficient general-purpose model",
}


# ----------------------------------------------------------------------
# SIDEBAR - SYSTEM STATUS & CONTROLS
# ----------------------------------------------------------------------
with st.sidebar:
    st.header("System Status")
    
    # Health check
    if st.button("Check Health"):
        try:
            with st.spinner("Checking health..."):
                response = requests.get(HEALTH_ENDPOINT, timeout=10)
                if response.ok:
                    st.session_state.health_status = response.json()
                    st.success("Online")
                else:
                    st.error("System Issues")
        except Exception as e:
            st.error(f"Cannot connect to API: {e}")
    
    # Display cached health status
    if st.session_state.health_status:
        with st.expander("System Details"):
            st.json(st.session_state.health_status)

    st.header("Collections")
    
    if st.button("Load Collections"):
        try:
            # Load collections from both sources
            chromadb_collections = get_chromadb_collections()
            
            # Combine and deduplicate
            all_collections = list(set(chromadb_collections))
            st.session_state.collections = all_collections
            st.success("Collections loaded!")
        except Exception as e:
            st.error(f"Error: {e}")
    
    # Display collections
    collections = st.session_state.collections
    if collections:
        for collection in collections:
            st.text(f"{collection}")
    else:
        st.info("Click 'Load Collections' to see available databases")

# Get collections for main interface
try:
    if not st.session_state.collections:
        chromadb_collections = get_chromadb_collections()
        collections = list(set(chromadb_collections))
        st.session_state.collections = collections
    else:
        collections = st.session_state.collections
except:
    collections = []

# ----------------------------------------------------------------------
# MAIN INTERFACE
# ----------------------------------------------------------------------
# Chat mode selection
chat_mode = st.radio(
    "Select Mode:",
    ["Direct Chat", "AI Agent Simulation", "Create Agent", "Document Generator", "Session History"],
    horizontal=True
)


# ----------------------------------------------------------------------
# DIRECT CHAT MODE
# ----------------------------------------------------------------------
if chat_mode == "Direct Chat":
    st.markdown("---")
    
    # Create tabs for chat functionality
    chat_tab, upload_tab = st.tabs(["Chat Interface", "Document Upload"])
    
    with chat_tab:
        # Model selection (common for all modes)
        col1, col2 = st.columns([2, 1])
        with col1:
            mode = st.selectbox("Select AI Model:", list(model_key_map.keys()))
            if model_key_map[mode] in model_descriptions:
                st.info(model_descriptions[model_key_map[mode]])
                
        # RAG Configuration
        use_rag = st.checkbox("Use RAG (Retrieval Augmented Generation)")
        if use_rag and collections:
            collection_name = st.selectbox("Document Collection:", collections)
        else:
            collection_name = None
            
        if use_rag and not collections:
            st.warning("No collections available. Upload documents first or check your ChromaDB connection.")
        
        # Chat Input
        user_input = st.text_area(
            "Ask your question:", 
            height=100, 
            placeholder="Example: Analyze this document to determine if it meets our quality standards..."
        )

        if st.button("Get Analysis", type="primary"):
            if not user_input:
                st.warning("Please enter a question.")
            elif use_rag and not collection_name:
                st.error("Please select a collection for RAG mode.")
            else:
                payload = {
                    "query": user_input,
                    "model": model_key_map[mode],
                    "use_rag": use_rag,
                    "collection_name": collection_name if use_rag else None
                }

                with st.spinner(f"{mode} is analyzing your question..."):
                    status_placeholder = st.empty()
                    
                    try:
                        status_placeholder.info("Connecting to AI model...")
                        
                        response = requests.post(CHAT_ENDPOINT, json=payload, timeout=300)  
                        
                        if response.ok:
                            result = response.json().get("response", "")
                            status_placeholder.empty()
                            
                            st.success("Analysis Complete:")
                            st.markdown("### Analysis Results:")
                            st.markdown(result)
                            
                            if "response_time_ms" in response.json():
                                response_time = response.json()["response_time_ms"]
                                st.caption(f"Response time: {response_time/1000:.2f} seconds")
                                
                        else:
                            status_placeholder.empty()
                            error_detail = response.json().get("detail", response.text) if response.headers.get("content-type") == "application/json" else response.text
                            
                            if "model" in error_detail and "not found" in error_detail:
                                st.error("Model is loading for the first time. This may take 1-2 minutes. Please try again.")
                                st.info("Tip: The first request to each model takes longer as it loads into memory.")
                            else:
                                st.error(f"Error {response.status_code}: {error_detail}")
                                
                    except requests.exceptions.Timeout:
                        status_placeholder.empty()
                        st.error("Request timed out. The model might be loading - please try again in a moment.")
                        st.info("Large models can take 1-2 minutes to load on first use.")
                    except requests.exceptions.RequestException as e:
                        status_placeholder.empty()
                        st.error(f"Request failed: {e}")
        
        # Chat History Section
        st.markdown("---")
        st.header("Chat History")
        if st.button("Load Chat History"):
            try:
                with st.spinner("Loading chat history..."):
                    response = requests.get(HISTORY_ENDPOINT, timeout=10)
                    if response.ok:
                        history = response.json()
                        if not history:
                            st.info("No chat history found.")
                        else:
                            for record in reversed(history[-10:]):
                                with st.expander(f"💬 {record['timestamp'][:19]}"):
                                    st.markdown(f"**User:** {record['user_query']}")
                                    st.markdown(f"**Response:** {record['response']}")
                    else:
                        st.error(f"Failed to fetch history: {response.status_code}")
            except requests.exceptions.RequestException as e:
                st.error(f"Request failed: {e}")
    
    with upload_tab:
        collections = get_chromadb_collections()
        render_upload_component(
            available_collections= collections,
            load_collections_func= get_chromadb_collections,
            create_collection_func= create_collection,
            upload_endpoint=f"{CHROMADB_API}/documents/upload-and-process",
            job_status_endpoint=f"{CHROMADB_API}/jobs/{{job_id}}"
        )
        # st.header("Document Upload & Processing")
        # st.info("Upload documents to create or enhance your knowledge base for RAG-powered conversations.")
        
        # # Collection Management Section
        # st.subheader("Collection Management")
        
        # collection_action = st.radio(
        #     "Choose action:",
        #     ["Use Existing Collection", "Create New Collection"],
        #     horizontal=True
        # )
        
        # if collection_action == "Use Existing Collection":
        #     if collections:
        #         target_collection = st.selectbox(
        #             "Select collection to add documents:",
        #             collections,
        #             help="Choose an existing collection to add your documents"
        #         )
        #     else:
        #         st.warning("No collections available. Create a new collection first.")
        #         target_collection = None
        # else:
        #     # Create new collection
        #     new_collection_name = st.text_input(
        #         "New collection name:",
        #         placeholder="e.g., legal-documents, contracts-2024, policies",
        #         help="Enter a descriptive name for your new collection"
        #     )
            
        #     if new_collection_name:
        #         if st.button("Create Collection", type="secondary"):
        #             try:
        #                 with st.spinner("Creating collection..."):
        #                     response = create_collection(new_collection_name)
                            
        #                     if response.status_code == 200:
        #                         st.success(f"Collection '{new_collection_name}' created successfully!")
        #                         # Refresh collections list
        #                         st.session_state.collections.append(new_collection_name)
        #                         target_collection = new_collection_name
        #                     else:
        #                         error_detail = response.json().get("detail", response.text) if response.headers.get("content-type") == "application/json" else response.text
        #                         st.error(f"Failed to create collection: {error_detail}")
        #                         target_collection = None
                                
        #             except Exception as e:
        #                 st.error(f"Error creating collection: {e}")
        #                 target_collection = None
                        
        #         target_collection = new_collection_name if new_collection_name else None
        #     else:
        #         target_collection = None
        
        # # Document Upload Section
        # if target_collection:
        #     st.subheader("Upload Documents")
            
        #     # File uploader
        #     uploaded_files = st.file_uploader(
        #         "Choose files to upload",
        #         type=['pdf', 'docx', 'txt', 'xlsx', 'pptx', 'html', 'csv'],
        #         accept_multiple_files=True,
        #         help="Supported formats: PDF, Word, Text, Excel, PowerPoint, HTML, CSV"
        #     )
            
        #     if uploaded_files:
        #         st.write(f"**Selected files ({len(uploaded_files)}):**")
        #         total_size = 0
        #         for file in uploaded_files:
        #             file_size = len(file.getvalue()) / 1024 / 1024  # Size in MB
        #             total_size += file_size
        #             st.write(f"• {file.name} ({file_size:.2f} MB)")
                
        #         st.info(f"Total size: {total_size:.2f} MB")
                
        #         # Processing Options
        #         with st.expander("Processing Options", expanded=False):
        #             col1, col2 = st.columns(2)
                    
        #             with col1:
        #                 chunk_size = st.number_input(
        #                     "Chunk Size", 
        #                     min_value=500, 
        #                     max_value=2000, 
        #                     value=1000,
        #                     help="Size of text chunks for processing"
        #                 )
                        
        #                 chunk_overlap = st.number_input(
        #                     "Chunk Overlap", 
        #                     min_value=50, 
        #                     max_value=500, 
        #                     value=200,
        #                     help="Overlap between consecutive chunks"
        #                 )
                    
        #             with col2:
        #                 use_openai_vision = st.checkbox(
        #                     "Use OpenAI Vision", 
        #                     value=True,
        #                     help="Enhanced image analysis using OpenAI"
        #                 )
                        
        #                 openai_key = st.text_input(
        #                     "OpenAI API Key (optional)",
        #                     type="password",
        #                     help="Required for OpenAI Vision features"
        #                 ) if use_openai_vision else None
                        
        #                 store_images = st.checkbox(
        #                     "Store Images", 
        #                     value=True,
        #                     help="Extract and store images from documents"
        #                 )
                
        #         # Upload and Process Button
        #         if st.button("Upload and Process Documents", type="primary"):
        #             if not target_collection:
        #                 st.error("Please select or create a collection first.")
        #             else:
        #                 try:
        #                     # 1) Kick off the job
        #                     status_text = st.empty()
        #                     progress_bar = st.progress(0)
        #                     status_text.text("Submitting job…")
        #                     resp = upload_documents_to_chromadb(
        #                         uploaded_files,
        #                         target_collection,
        #                         openai_key if use_openai_vision else None
        #                     )
        #                     resp.raise_for_status()
        #                     job_id = resp.json()["job_id"]
        #                     status_text.text(f"Job ID: {job_id} – waiting for progress…")
                            
        #                     # 2) Poll the /jobs endpoint until complete
        #                     while True:
        #                         stat = requests.get(f"{CHROMADB_URL}/jobs/{job_id}", timeout=5).json()
        #                         state = stat.get("status", "unknown")
        #                         if state == "running":
        #                             done = stat.get("processed_chunks", 0)
        #                             total = stat.get("total_chunks",   1)
        #                             progress = min(1.0, done/total)
        #                             status_text.text(f"Processed {done}/{total} chunks…")
        #                             progress_bar.progress(progress)
        #                             time.sleep(1)
        #                             continue
        #                         elif state == "success":
        #                             status_text.success("All chunks ingested!")
        #                             progress_bar.progress(1.0)
        #                         else:
        #                             status_text.error(f"Job failed: {state}")
        #                         break
                            
        #                 except Exception as e:
        #                     progress_bar.progress(0)
        #                     status_text.error(f"Upload failed: {e}")
        #                     raise

        #                     # Helpful error messages
        #                     if "timeout" in str(e).lower():
        #                         st.info("Large files may take several minutes to process. Consider uploading fewer files at once.")
        #                     elif "connection" in str(e).lower():
        #                         st.info("Check that your ChromaDB service is running and accessible.")
        
        # else:
        #     st.warning("Please select an existing collection or create a new one to upload documents.")
        
        # Current Collections Status
        st.markdown("---")
        st.subheader("Current Collections")
        
        if st.button("Refresh Collections"):
            try:
                chromadb_collections = get_chromadb_collections()
                st.session_state.collections = chromadb_collections
                st.success(f"Found {len(chromadb_collections)} collections")
            except Exception as e:
                st.error(f"Error refreshing collections: {e}")
        
        if collections:
            for i, collection in enumerate(collections, 1):
                st.write(f"{i}. **{collection}**")
        else:
            st.info("No collections found. Create your first collection to get started!")
        

# ----------------------------------------------------------------------
# AI AGENT SIMULATION MODE
# ----------------------------------------------------------------------
elif chat_mode == "AI Agent Simulation":
    st.markdown("---")
    
    # Load agents with enhanced error handling
    col1_load, col2_load = st.columns([1, 2])
    
    with col1_load:
        if st.button("Refresh Agent List"):
            try:
                with st.spinner("Loading agents from database..."):
                    agents_response = requests.get(f"{FASTAPI_API}/get-agents", timeout=10)
                    if agents_response.status_code == 200:
                        agent_data = agents_response.json()
                        st.session_state.agents_data = agent_data.get("agents", [])
                        st.success(f"Loaded {len(st.session_state.agents_data)} agents")
                    else:
                        st.warning(f"Could not load agents (Status: {agents_response.status_code})")
                        st.error(f"Response: {agents_response.text}")
            except Exception as e:
                st.error(f"Error loading agents: {e}")
    
    with col2_load:
        if st.session_state.agents_data:
            total_agents = len(st.session_state.agents_data)
            active_agents = sum(1 for agent in st.session_state.agents_data if agent.get("is_active", True))
            st.metric("Total Agents", total_agents, delta=f"{active_agents} active")

    # Check if we have agents data
    if st.session_state.agents_data:
        agents = st.session_state.agents_data
        
        # Create agent choices dictionary for selection
        agent_choices = {f"{agent['name']} ({agent['model_name']})": agent["id"] for agent in agents}
        
        # Enhanced agents display
        agents_table_data = []
        for agent in agents:
            agents_table_data.append({
                "ID": agent.get("id", "N/A"),
                "Name": agent.get("name", "Unknown"),
                "Model": agent.get("model_name", "Unknown"),
                "Queries": agent.get("total_queries", 0),
                "Avg Response": f"{agent.get('avg_response_time_ms', 0):.0f}ms" if agent.get('avg_response_time_ms') else "N/A",
                "Success Rate": f"{agent.get('success_rate', 0)*100:.1f}%" if agent.get('success_rate') else "N/A",
                "Status": "Active" if agent.get("is_active", True) else "Inactive",
                "Created": agent.get("created_at", "Unknown")[:10] if agent.get("created_at") else "Unknown",
                "System Prompt": agent.get("system_prompt", "")[:100] + ("..." if len(agent.get("system_prompt", "")) > 100 else ""),
                "User Template": agent.get("user_prompt_template", "")[:100] + ("..." if len(agent.get("user_prompt_template", "")) > 100 else "")
            })
        
        # Display agents table
        st.dataframe(agents_table_data, use_container_width=True, height=400)
        
        # Single Agent Analysis Section
        st.subheader("Single Agent Analysis")
        
        # Analysis content
        analysis_content = st.text_area(
            "Content for Agent Analysis", 
            placeholder="Paste contract text,documents, or content for analysis...",
            height=150
        )
        
        selected_agents = st.multiselect(
            "Select Specialized Agents for Analysis", 
            list(agent_choices.keys()),
            help="Choose multiple agents to get different perspectives on your content"
        )

        # Analysis type selection
        analysis_type = st.selectbox(
            "Analysis Type:",
            ["Direct Analysis", "RAG-Enhanced Analysis"],
            help="Direct uses content as-is, RAG adds context from your knowledge base"
        )

        if analysis_type == "RAG-Enhanced Analysis":
            if collections:
                collection_name = st.selectbox("Collection to utilize Retrieval Augmented Generation:", collections)
            else:
                st.warning("No collections available for RAG analysis")
                collection_name = None
        else:
            collection_name = None

        if st.button("Run Agent Analysis", type="primary"):
            if not analysis_content or not selected_agents:
                st.warning("Please provide content and select at least one agent.")
            else:
                agent_ids = [agent_choices[name] for name in selected_agents]
                
                # Choose endpoint based on analysis type
                if analysis_type == "RAG-Enhanced Analysis" and collection_name:
                    payload = {
                        "query_text": analysis_content,
                        "collection_name": collection_name,
                        "agent_ids": agent_ids
                    }
                    endpoint = f"{FASTAPI_API}/rag-check"
                else:
                    payload = {
                        "data_sample": analysis_content,
                        "agent_ids": agent_ids
                    }
                    endpoint = f"{FASTAPI_API}/compliance-check"
                
                with st.spinner("Specialized agents are analyzing the content..."):
                    status_placeholder = st.empty()
                    try:
                        status_placeholder.info("Connecting to AI model...")
                        response = requests.post(endpoint, json=payload, timeout=300)
                        
                        if response.status_code == 200:
                            result = response.json()
                            
                            # Display results
                            agent_responses = result.get("agent_responses", {})
                            if agent_responses:
                                for agent_name, analysis in agent_responses.items():
                                    with st.expander(f"{agent_name} Analysis", expanded=True):
                                        st.markdown(analysis)
                            else:
                                # Handle compliance check format
                                details = result.get("details", {})
                                for idx, analysis in details.items():
                                    agent_name = analysis.get("agent_name", f"Agent {idx}")
                                    reason = analysis.get("reason", analysis.get("raw_text", "No analysis"))
                                    
                                    with st.expander(f"{agent_name} Analysis", expanded=True):
                                        st.markdown(reason)
                                        
                            if "session_id" in result:
                                st.info(f"Analysis Session ID: {result['session_id']}")
                                
                        else:
                            error_detail = response.json().get("detail", response.text) if response.headers.get("content-type") == "application/json" else response.text
                            st.error(f"Error {response.status_code}: {error_detail}")
                            
                    except requests.exceptions.Timeout:
                        status_placeholder.empty()
                        st.error("Request timed out. The model might be loading - please try again in a moment.")
                        st.info("Large models can take 1-2 minutes to load on first use.")
                        # st.error("Analysis timed out. Try with shorter content or fewer agents.")
                    except Exception as e:
                        status_placeholder.empty()
                        st.error(f"Analysis failed: {str(e)}")

        st.markdown("---")
        
        # Multi-Agent Debate Sequence Section
        st.subheader("Multi-Agent Debate Sequence")
        st.info("**How it works**: Create a sequence of agents that will debate in order. Each agent uses their pre-configured LLM and prompts.")
        
        # Initialize session state for agent sequence
        if "debate_sequence" not in st.session_state:
            st.session_state["debate_sequence"] = []
        
        # Agent sequence builder
        col1, col2 = st.columns([3, 1])
        with col1:
            new_agent_to_add = st.selectbox(
                "Add Agent to Debate Sequence", 
                ["--Select an Agent--"] + list(agent_choices.keys()), 
                key="debate_sequence_select"
            )
        with col2:
            if st.button("+ Add to Sequence", key="add_agent_debate"):
                if new_agent_to_add != "--Select an Agent--" and new_agent_to_add not in st.session_state["debate_sequence"]:
                    st.session_state["debate_sequence"].append(new_agent_to_add)
                    st.success(f"Added {new_agent_to_add}")
                    st.rerun()
                elif new_agent_to_add in st.session_state["debate_sequence"]:
                    st.warning("Agent already in sequence!")
            
            if st.button("Clear All Sequences"):
                st.session_state["debate_sequence"] = []
                st.rerun()
        
        # Display current debate sequence
        if st.session_state["debate_sequence"]:
            st.write("**Current Debate Sequence:**")
            for i, agent_name in enumerate(st.session_state["debate_sequence"], 1):
                col1, col2, col3 = st.columns([1, 4, 1])
                with col1:
                    st.write(f"**{i}.**")
                with col2:
                    # Show agent info
                    agent_data = next((agent for agent in agents if f"{agent['name']} ({agent['model_name']})" == agent_name), None)
                    if agent_data:
                        st.write(f"**{agent_data['name']}** using *{agent_data['model_name']}*")
                    else:
                        st.write(f"{agent_name}")
                with col3:
                    if st.button("Remove", key=f"remove_{i}", help="Remove from sequence"):
                        st.session_state["debate_sequence"].remove(agent_name)
                        st.rerun()
            
            st.markdown("---")
        else:
            st.info("Add agents to create a debate sequence")
        
        # Content and collection selection for debate
        debate_content = st.text_area(
            "Content for Multi-Agent Debate", 
            placeholder="Enter the content that agents will debate about...",
            height=120, 
            key="debate_content"
        )
        
        # RAG selection for debate
        use_rag_debate = st.checkbox("Use RAG for Debate Context", key="rag_debate")
        if use_rag_debate and collections:
            collection_for_debate = st.selectbox(
                "ChromaDB Collection (for RAG context):", 
                collections,
                help="Select a collection to provide context for the debate"
            )
        else:
            if use_rag_debate:
                st.warning("No collections available. Agents will debate without RAG context.")
            collection_for_debate = None
        
        # Start debate button
        if st.button("Start Multi-Agent Debate", type="primary", key="start_debate"):
            if not debate_content:
                st.warning("Please provide content for the agents to debate.")
            elif not st.session_state["debate_sequence"]:
                st.warning("Please add at least one agent to the debate sequence.")
            else:
                # Prepare the debate
                sequence_agent_ids = []
                for agent_name in st.session_state["debate_sequence"]:
                    agent_id = agent_choices.get(agent_name)
                    if agent_id:
                        sequence_agent_ids.append(agent_id)
                
                if not sequence_agent_ids:
                    st.error("Unable to find agent IDs. Please reload agents and try again.")
                else:
                    # Show what's about to happen
                    with st.expander("Debate Setup", expanded=True):
                        st.write(f"**Content**: {debate_content[:100]}{'...' if len(debate_content) > 100 else ''}")
                        st.write(f"**Agents in sequence**: {len(sequence_agent_ids)}")
                        st.write(f"**Using RAG**: {'Yes' if use_rag_debate and collection_for_debate else 'No'}")
                        if use_rag_debate and collection_for_debate:
                            st.write(f"**Collection**: {collection_for_debate}")
                    
                    # Prepare payload based on whether we're using RAG or not
                    if use_rag_debate and collection_for_debate:
                        debate_payload = {
                            "query_text": debate_content,
                            "collection_name": collection_for_debate,
                            "agent_ids": sequence_agent_ids
                        }
                        endpoint = f"{FASTAPI_API}/rag-debate-sequence"
                    else:
                        debate_payload = {
                            "data_sample": debate_content,
                            "agent_ids": sequence_agent_ids
                        }
                        endpoint = f"{FASTAPI_API}/compliance-check"  # Use compliance check for non-RAG debate
                    
                    # Start the debate
                    with st.spinner(f"{len(sequence_agent_ids)} agents are debating..."):
                        status_placeholder = st.empty()
                        try:
                            status_placeholder.info("Connecting to debate service...")
                            response = requests.post(endpoint, json=debate_payload, timeout=300)
                            
                            if response.status_code == 200:
                                result = response.json()
                                st.success("Multi-Agent Debate Complete!")
                                
                                # Display results
                                session_id = result.get("session_id")
                                if session_id:
                                    st.info(f"Debate Session ID: `{session_id}`")
                                
                                # Show debate results
                                if "debate_chain" in result:
                                    debate_chain = result["debate_chain"]
                                    st.subheader("Debate Sequence Results")
                                    
                                    for i, round_result in enumerate(debate_chain, 1):
                                        agent_name = round_result.get('agent_name', 'Unknown Agent')
                                        response_text = round_result.get("response", "No response")
                                        
                                        with st.expander(f"Round {i}: {agent_name}", expanded=i<=2):
                                            st.markdown(response_text)
                                            
                                            # Show metadata if available
                                            if "agent_id" in round_result:
                                                st.caption(f"Agent ID: {round_result['agent_id']}")
                                
                                elif "details" in result:
                                    # Handle compliance check format
                                    st.subheader("Agent Analysis Results")
                                    details = result["details"]
                                    
                                    for idx, analysis in details.items():
                                        agent_name = analysis.get("agent_name", f"Agent {idx}")
                                        reason = analysis.get("reason", analysis.get("raw_text", "No analysis"))
                                        
                                        with st.expander(f"{agent_name}", expanded=True):
                                            st.markdown(reason)
                                
                                elif "agent_responses" in result:
                                    # Handle agent_responses format
                                    st.subheader("Agent Analysis Results")
                                    agent_responses = result["agent_responses"]
                                    
                                    for agent_name, response_text in agent_responses.items():
                                        with st.expander(f"{agent_name}", expanded=True):
                                            st.markdown(response_text)
                                
                                else:
                                    st.json(result)  # Fallback to show raw result
                                    
                            else:
                                error_detail = response.json().get("detail", response.text) if response.headers.get("content-type") == "application/json" else response.text
                                st.error(f"Error {response.status_code}: {error_detail}")
                                
                        except requests.exceptions.Timeout:
                            status_placeholder.empty()
                            st.error("Request timed out. The model might be loading - please try again in a moment.")
                            st.info("Large models can take 1-2 minutes to load on first use.")
                        except Exception as e:
                            status_placeholder.empty()
                            st.error(f"Debate failed: {str(e)}")

        # Help section
        with st.expander("How Multi-Agent Debate Works"):
            st.markdown("""
            **Multi-Agent Debate Process:**
            
            1. **Create Sequence**: Add agents in the order you want them to participate
            2. **Each agent uses their own configured LLM** (LLama3)
            3. **Each agent applies their specialized prompts** (Systems Engineering analysis, Quality Control Engineer, etc.)
            4. **Sequential Analysis**: Agents review the content in order
            5. **Optional RAG**: Choose a collection to provide additional context
            
            **Example Sequence:**
            - Risk Analyzer (Llama3) → Identifies risks
            - Systems Engineer (LLama3) → Reviews requirements
            - Test Engineer (Llama3) → Provides test context
            
            **Why this is better than selecting models again:**
            - Agents already have optimized LLM+prompt combinations
            - Each agent brings their specialized expertise
            - No need to duplicate model selection
            """)

    else:
        # No agents available
        st.info("Click 'Refresh Agent List' to load your specialized agents")
        st.markdown("""
        **No agents found!** 
        
        To use Agent Simulation mode:
        1. Go to the 'Create Agent' tab
        2. Create some specialized agents
        3. Come back here to simulate multi-agent analysis
        """)

    # Footer
    st.markdown("---")
    st.markdown("* All agents provide analysis for informational purposes only*")

# ----------------------------------------------------------------------
# CREATE AGENT MODE (WITH MANAGEMENT SUB-MODES)
# ----------------------------------------------------------------------
elif chat_mode == "Create Agent":
    st.markdown("---")
    st.header("Agent Management")
    
    # Agent management sub-modes
    agent_mode = st.radio(
        "Select Action:",
        ["Create New Agent", "Manage Existing Agents"],
        horizontal=True
    )
    
    # ----------------------------------------------------------------------
    # CREATE NEW AGENT SUB-MODE
    # ----------------------------------------------------------------------
    if agent_mode == "Create New Agent":
        st.subheader("Create a New Agent")
        
        # Enhanced agent templates (keeping your existing templates)
        enhanced_templates = {
            "Systems Engineering Agent": {
                "description": "Systems engineering specialist applying SEBoK (Systems Engineering Body of Knowledge) principles to review and enhance requirement development processes",
                "system_prompt": """You are an experienced systems engineer with 15+ years of experience in complex system development across aerospace, defense, and technology sectors. Your expertise spans the full systems engineering lifecycle with deep knowledge of SEBoK principles. Your role is to:

                1. **Requirements Analysis**: Evaluate requirements for completeness, consistency, clarity, and traceability
                2. **SEBoK Application**: Apply Systems Engineering Body of Knowledge best practices and standards
                3. **Stakeholder Assessment**: Analyze stakeholder needs translation into verifiable requirements
                4. **Architecture Alignment**: Ensure requirements support system architecture and design decisions
                5. **Verification Planning**: Assess testability and verification approaches for each requirement
                6. **Risk Identification**: Identify requirement-related risks and mitigation strategies
                7. **Process Improvement**: Recommend enhancements to requirement development processes

                **SEBoK Framework Application:**
                - Apply life cycle processes (ISO/IEC/IEEE 15288)
                - Utilize requirements engineering best practices
                - Ensure stakeholder needs are properly captured and managed
                - Maintain traceability throughout system hierarchy
                - Consider system context and boundary definitions
                - Apply appropriate requirement attributes and characteristics

                **Technical Approach:**
                - Evaluate requirements against SMART criteria (Specific, Measurable, Achievable, Relevant, Time-bound)
                - Assess requirement quality attributes (unambiguous, complete, consistent, testable)
                - Review for proper categorization (functional, performance, interface, constraint)
                - Analyze requirement dependencies and interactions
                - Consider system integration and interface requirements
                - Evaluate for proper abstraction levels across system hierarchy""",

                "user_prompt": """As a systems engineering specialist, review the following requirement development artifact and provide a comprehensive SEBoK-based assessment:

                {data_sample}

                Please provide a detailed systems engineering analysis:

                1. **Requirements Quality Assessment** (completeness, clarity, consistency, traceability evaluation)
                2. **SEBoK Compliance Review** (alignment with systems engineering best practices and standards)
                3. **Stakeholder Analysis** (adequacy of stakeholder need capture and translation)
                4. **Architecture Considerations** (requirement support for system design and interfaces)
                5. **Verification & Validation Planning** (testability assessment and V&V approach recommendations)
                6. **Risk Analysis** (requirement-related risks and technical concerns)
                7. **Process Recommendations** (improvements to requirement development methodology)
                8. **Traceability Assessment** (requirement hierarchy and dependency analysis)
                9. **Next Steps** (prioritized actions to enhance requirement quality and process)

                Focus on applying rigorous systems engineering principles while providing practical, actionable recommendations for improvement.""",

                "temperature": 0.3,
                "max_tokens": 2000
            },
            
            "Test Engineering Agent": {
                "description": "Expert system test engineer specializing in comprehensive test case development, test plan review, and verification strategy across complex integrated systems",
                "system_prompt": """You are a senior system test engineer with 12+ years of experience in developing and executing test strategies for complex integrated systems across aerospace, automotive, telecommunications, and enterprise software domains. Your expertise encompasses the full testing lifecycle from planning through execution and reporting. Your role is to:

                1. **Test Strategy Development**: Design comprehensive test approaches covering functional, non-functional, and integration testing
                2. **Test Case Design**: Create detailed, traceable test cases using systematic design techniques (equivalence partitioning, boundary value analysis, state-based testing)
                3. **Test Plan Review**: Evaluate test plans for completeness, feasibility, risk coverage, and alignment with requirements
                4. **Verification Planning**: Develop verification strategies that demonstrate system compliance with specifications
                5. **Test Environment Design**: Specify test infrastructure, tooling, and data requirements
                6. **Risk-Based Testing**: Prioritize testing efforts based on system criticality and failure impact analysis
                7. **Test Automation Strategy**: Identify automation opportunities and develop sustainable test frameworks

                **Testing Framework Application:**
                - Apply IEEE 829 test documentation standards
                - Utilize ISO/IEC/IEEE 29119 testing principles and processes
                - Implement systematic test design techniques and coverage criteria
                - Ensure bidirectional traceability between requirements and test cases
                - Apply risk-based testing methodologies (ISO 31000)
                - Consider system integration testing approaches (big-bang, incremental, sandwich)
                - Evaluate test completion criteria and exit conditions

                **Technical Approach:**
                - Design test cases covering positive, negative, and boundary conditions
                - Develop test scenarios for system-level behaviors and emergent properties
                - Create test procedures with clear setup, execution, and validation steps
                - Specify test data requirements and management strategies
                - Define test environment configurations and dependencies
                - Plan for defect management and regression testing cycles
                - Consider performance, security, usability, and reliability testing aspects""",
                    
                    "user_prompt": """As a senior system test engineer, analyze the following testing artifact and provide a comprehensive assessment:

                {data_sample}

                Please provide a detailed system testing analysis:

                1. **Test Strategy Assessment** (completeness of testing approach and methodology)
                2. **Test Case Quality Review** (clarity, traceability, executability, and coverage evaluation)
                3. **Requirements Coverage Analysis** (mapping between requirements and test cases, gap identification)
                4. **Test Plan Evaluation** (feasibility, resource allocation, timeline, and risk considerations)
                5. **Test Environment & Infrastructure** (adequacy of test setup, tooling, and data requirements)
                6. **Integration Testing Strategy** (system integration approach and interface testing coverage)
                7. **Non-Functional Testing Coverage** (performance, security, reliability, usability considerations)
                8. **Risk Analysis & Mitigation** (testing risks and contingency planning)
                9. **Test Automation Opportunities** (automation feasibility and ROI assessment)
                10. **Process Recommendations** (improvements to testing methodology and practices)
                11. **Next Steps** (prioritized actions to enhance test quality and execution)

                Focus on providing practical, actionable recommendations that improve test effectiveness, efficiency, and system quality assurance.""",
                    
                    "temperature": 0.2,
                    "max_tokens": 2200
            },
            "Quality Engineering Agent": {
                "description": "Expert quality engineering specialist focused on implementing comprehensive quality management systems, process improvement, and quality assurance across the entire product lifecycle",
                "system_prompt": """You are a senior quality engineering professional with 15+ years of experience implementing quality management systems across manufacturing, software development, aerospace, medical devices, and automotive industries. Your expertise spans quality planning, process control, continuous improvement, and regulatory compliance. Your role is to:

                1. **Quality Management Systems**: Design and implement QMS frameworks (ISO 9001, AS9100, ISO 13485, IATF 16949)
                2. **Process Quality Control**: Establish statistical process control, quality metrics, and performance monitoring
                3. **Quality Planning**: Develop quality plans, control plans, and quality gates throughout product lifecycle
                4. **Risk Management**: Implement quality risk assessment methodologies (FMEA, FTA, Risk Priority Numbers)
                5. **Continuous Improvement**: Lead quality improvement initiatives using Lean Six Sigma, DMAIC, and Kaizen methodologies
                6. **Supplier Quality**: Establish supplier quality requirements, audits, and performance management
                7. **Regulatory Compliance**: Ensure adherence to industry standards and regulatory requirements
                8. **Quality Analytics**: Develop quality dashboards, trend analysis, and predictive quality models

                **Quality Framework Application:**
                - Apply Total Quality Management (TQM) principles
                - Implement Plan-Do-Check-Act (PDCA) cycles for continuous improvement
                - Utilize Statistical Quality Control (SQC) and Design of Experiments (DOE)
                - Apply quality cost models (Prevention, Appraisal, Internal/External Failure costs)
                - Implement configuration management and change control processes
                - Ensure traceability and documentation control throughout lifecycle
                - Apply quality gate criteria and stage-gate reviews

                **Technical Approach:**
                - Establish quality objectives with measurable KPIs and targets
                - Design quality control checkpoints and inspection strategies
                - Implement corrective and preventive action (CAPA) processes
                - Develop quality training and competency management programs
                - Create quality documentation hierarchies and control systems
                - Establish customer feedback loops and satisfaction measurement
                - Design quality audit programs and management review processes
                - Implement quality culture transformation and organizational change management""",
                    
                    "user_prompt": """As a senior quality engineering specialist, analyze the following quality-related artifact and provide a comprehensive assessment:

                {data_sample}

                Please provide a detailed quality engineering analysis:

                1. **Quality System Assessment** (QMS framework evaluation and compliance review)
                2. **Process Quality Analysis** (process capability, control measures, and statistical analysis)
                3. **Quality Planning Evaluation** (quality objectives, control plans, and gate criteria)
                4. **Risk Assessment Review** (quality risks identification, FMEA analysis, and mitigation strategies)
                5. **Metrics & KPI Analysis** (quality measurement systems and performance indicators)
                6. **Continuous Improvement Opportunities** (improvement initiatives and optimization recommendations)
                7. **Compliance & Standards Review** (regulatory requirements and industry standard adherence)
                8. **Cost of Quality Analysis** (prevention, appraisal, and failure cost assessment)
                9. **Supplier Quality Considerations** (supply chain quality requirements and management)
                10. **Quality Culture & Training** (organizational quality maturity and competency gaps)
                11. **Documentation & Traceability** (quality record management and audit trail adequacy)
                12. **Action Plan Development** (prioritized quality improvement roadmap and implementation strategy)

                Focus on providing systematic, data-driven recommendations that enhance quality performance, reduce defects, improve customer satisfaction, and drive organizational quality maturity.""",
                    
                    "temperature": 0.2,
                    "max_tokens": 2400
            }
        }

        # Enhanced layout with better organization (keeping your existing layout)
        col1, col2 = st.columns([1, 1])

        with col1:
            st.subheader("Basic Configuration")
            
            # Agent name with validation
            agent_name = st.text_input(
                "Agent Name", 
                placeholder="e.g., Systems Engineer Agent, Systems Test Agent, QC Agent",
                help="Choose a descriptive name that reflects the agent's specialty"
            )
            
            # Real-time validation feedback
            if agent_name:
                if len(agent_name) < 3:
                    st.warning("Agent name should be at least 3 characters")
                elif len(agent_name) > 100:
                    st.warning("Agent name should be less than 100 characters")
                else:
                    st.success("Agent name looks good")
            
            # Enhanced agent type selection with descriptions
            agent_type = st.selectbox(
                "Agent Type Template", 
                ["Custom"] + list(enhanced_templates.keys()),
                help="Select a predefined template optimized for specific tasks"
            )
            
            # Display template description
            if agent_type != "Custom":
                template = enhanced_templates[agent_type]
                st.info(f"**{agent_type}**: {template['description']}")

            # Enhanced model selection with status indicators
            st.subheader("Model Selection")
            
            # Get available models with status
            available_models = st.session_state.available_models or get_available_models_cached()
            
            if available_models:
                # Display model options with enhanced info
                model_options = []
                model_status = {}
                
                for model in available_models:
                    status = check_model_status(model)
                    if status == "available":
                        model_options.append(f"{model}")
                        model_status[model] = "available"
                    else:
                        model_options.append(f"{model}")
                        model_status[model] = "degraded"
                
                selected_model_display = st.selectbox("Select Model", model_options)
                split_parts = selected_model_display.split(" ", 1)
                agent_model_name = split_parts[1] if len(split_parts) > 1 else split_parts[0]
                
                # Display model information
                if agent_model_name in model_descriptions:
                    st.info(model_descriptions[agent_model_name])
                
                # Model status info
                status = model_status.get(agent_model_name, "unknown")
                if status == "available":
                    st.success("Model is loaded and ready")
                else:
                    st.error("Model is not available")
                    
            else:
                st.error("No models available. Please check your model configuration.")
                agent_model_name = None

        with col2:
            st.subheader("Agent Prompts & Configuration")
            
            # Auto-populate from template
            if agent_type != "Custom":
                template = enhanced_templates[agent_type]
                default_system_prompt = template["system_prompt"]
                default_user_prompt = template["user_prompt"]
                default_temperature = template.get("temperature", 0.7)
                default_max_tokens = template.get("max_tokens", 1000)
            else:
                default_system_prompt = ""
                default_user_prompt = ""
                default_temperature = 0.7
                default_max_tokens = 1000

            # System prompt with enhanced validation
            system_prompt = st.text_area(
                "System Prompt (Define Agent's Role & Expertise)",
                value=default_system_prompt,
                height=250,
                help="Define the agent's role, expertise, analysis approach, and output format",
                placeholder="You are an expert specializing in..."
            )
            
            # Character count and validation for system prompt
            if system_prompt:
                char_count = len(system_prompt)
                if char_count < 100:
                    st.warning(f"System prompt is quite short ({char_count} chars). Consider adding more detail.")
                elif char_count > 3000:
                    st.warning(f"System prompt is very long ({char_count} chars). Consider condensing.")
                else:
                    st.success(f"System prompt length is good ({char_count} chars)")

            # User prompt template with validation
            user_prompt_template = st.text_area(
                "User Prompt Template",
                value=default_user_prompt,
                height=150,
                help="Template for user queries. Must include {data_sample} placeholder for input content",
                placeholder="Analyze the following content:\n\n{data_sample}\n\nProvide detailed analysis..."
            )
            
            # Validation for user prompt template
            if user_prompt_template:
                if "{data_sample}" not in user_prompt_template:
                    st.error("User prompt template must include {data_sample} placeholder")
                else:
                    st.success("User prompt template is properly formatted")

            # Advanced configuration in expandable section
            with st.expander("Advanced Settings"):
                col1_adv, col2_adv = st.columns(2)
                
                with col1_adv:
                    temperature = st.slider(
                        "Temperature",
                        min_value=0.0,
                        max_value=1.0,
                        value=default_temperature,
                        step=0.1,
                        help="Controls creativity vs consistency. Lower = more focused and consistent responses"
                    )
                    
                    if temperature <= 0.3:
                        st.info("Low temperature: Highly consistent, focused responses")
                    elif temperature <= 0.7:
                        st.info("Medium temperature: Balanced creativity and consistency")
                    else:
                        st.info("High temperature: More creative and varied responses")
                
                with col2_adv:
                    max_tokens = st.number_input(
                        "Max Tokens",
                        min_value=100,
                        max_value=4000,
                        value=default_max_tokens,
                        step=100,
                        help="Maximum length of agent responses (roughly 1 token = 0.75 words)"
                    )
                    
                    estimated_words = int(max_tokens * 0.75)
                    st.caption(f"≈ {estimated_words} words maximum")

        # Enhanced validation and testing section (keeping your existing validation)
        # st.markdown("---")
        st.subheader("Validation & Testing")
        
        # Comprehensive validation
        validation_checks = {
            "Agent name provided": bool(agent_name and len(agent_name.strip()) >= 3),
            "Model selected": bool(agent_model_name),
            "System prompt substantial": bool(system_prompt and len(system_prompt.strip()) >= 100),
            "User prompt has {data_sample}": "{data_sample}" in user_prompt_template if user_prompt_template else False,
            "User prompt substantial": bool(user_prompt_template and len(user_prompt_template.strip()) >= 50),
            "No duplicate agent name": True  # You could add API check here
        }
        
        # Display validation results in columns
        col1_val, col2_val = st.columns(2)
        
        with col1_val:
            for check, status in list(validation_checks.items())[:3]:
                if status:
                    st.success(f"{check}")
                else:
                    st.error(f"{check}")
        
        with col2_val:
            for check, status in list(validation_checks.items())[3:]:
                if status:
                    st.success(f"{check}")
                else:
                    st.error(f"{check}")

        # Action buttons (keeping your existing create logic)
        col1_action, col2_action, col3_action = st.columns([1, 1, 1])
        
        
        with col2_action:
            if st.button("Create Agent", type="primary", use_container_width=True):
                if not all(validation_checks.values()):
                    st.warning("Please fix all validation errors before creating the agent")
                else:
                    # Enhanced agent payload
                    agent_payload = {
                        "name": agent_name.strip(),
                        "model_name": agent_model_name,
                        "system_prompt": system_prompt.strip(),
                        "user_prompt_template": user_prompt_template.strip()
                    }
                    
                    with st.spinner(f"Creating agent '{agent_name}'..."):
                        try:
                            response = requests.post(
                                f"{FASTAPI_API}/create-agent", 
                                json=agent_payload, 
                                timeout=30
                            )
                            
                            if response.status_code == 200:
                                result = response.json()
                                st.success(f"Agent '{agent_name}' created successfully!")
                                
                                # Display creation details
                                st.info(f"Agent ID: {result.get('agent_id', 'Unknown')}")
                                st.info(f"Model: {agent_model_name}")
                                st.info(f"Type: {agent_type}")
                                
                                # Clear session state to force refresh
                                st.session_state.agents_data = []
                                
                                # Quick action suggestions
                                st.markdown("**Next Steps:**")
                                st.markdown("- Switch to 'AI Agent Simulation' mode to test your new agent")
                                st.markdown("- Use 'Manage Existing Agents' tab to edit or manage agents")
                                st.markdown("- Create additional specialized agents for different tasks")
                                
                            else:
                                error_detail = response.json().get("detail", response.text) if response.headers.get("content-type") == "application/json" else response.text
                                
                                if "already exists" in error_detail.lower():
                                    st.error(f"Agent name '{agent_name}' already exists. Please choose a different name.")
                                else:
                                    st.error(f"Failed to create agent: {error_detail}")
                                    
                        except requests.exceptions.Timeout:
                            st.error("Request timed out. Please try again.")
                        except requests.exceptions.ConnectionError:
                            st.error("Cannot connect to the API. Please check if the server is running.")
                        except Exception as e:
                            st.error(f"Unexpected error: {str(e)}")

    # ----------------------------------------------------------------------
    # MANAGE EXISTING AGENTS SUB-MODE
    # ----------------------------------------------------------------------
    elif agent_mode == "Manage Existing Agents":
        st.subheader("Manage Existing Agents")
        
        # Load agents with enhanced error handling
        col1_load, col2_load = st.columns([1, 2])
        
        with col1_load:
            if st.button("Refresh Agent List", key="manage_refresh"):
                try:
                    with st.spinner("Loading agents from database..."):
                        agents_response = requests.get(f"{FASTAPI_API}/get-agents", timeout=10)
                        if agents_response.status_code == 200:
                            agent_data = agents_response.json()
                            st.session_state.agents_data = agent_data.get("agents", [])
                            st.success(f"Loaded {len(st.session_state.agents_data)} agents")
                        else:
                            st.warning(f"Could not load agents (Status: {agents_response.status_code})")
                except Exception as e:
                    st.error(f"Error loading agents: {e}")
        
        with col2_load:
            if st.session_state.agents_data:
                total_agents = len(st.session_state.agents_data)
                active_agents = sum(1 for agent in st.session_state.agents_data if agent.get("is_active", True))
                st.metric("Total Agents", total_agents, delta=f"{active_agents} active")

        # Enhanced agents display
        if st.session_state.agents_data:
            # Create enhanced table data
            agents_data = []
            for agent in st.session_state.agents_data:
                agents_data.append({
                    "ID": agent.get("id", "N/A"),
                    "Name": agent.get("name", "Unknown"),
                    "Model": agent.get("model_name", "Unknown"),
                    "Queries": agent.get("total_queries", 0),
                    "Status": "Active" if agent.get("is_active", True) else "Inactive",
                    "Created": agent.get("created_at", "Unknown")[:10] if agent.get("created_at") else "Unknown",
                    "System Prompt": agent.get("system_prompt", "")[:100] + ("..." if len(agent.get("system_prompt", "")) > 100 else ""),
                    "User Template": agent.get("user_prompt_template", "")[:100] + ("..." if len(agent.get("user_prompt_template", "")) > 100 else "")
                })
            
            # Display in a nice table
            st.dataframe(agents_data, use_container_width=True, height=400)
            
            # Management actions
            management_action = st.radio(
                "Management Action:",
                ["View Details", "Edit Agent", "Delete Agent"],
                horizontal=True,
                key="mgmt_action"
            )
            
            # Agent selection for all actions
            agent_options = [f"{agent['name']} (ID: {agent['id']})" for agent in st.session_state.agents_data]
            selected_agent_str = st.selectbox(
                "Select Agent:",
                ["--Select Agent--"] + agent_options,
                key="selected_agent_mgmt"
            )
            
            if selected_agent_str != "--Select Agent--":
                agent_id = int(selected_agent_str.split("ID: ")[1].rstrip(")"))
                selected_agent = next((agent for agent in st.session_state.agents_data if agent["id"] == agent_id), None)
                
                if selected_agent:
                    # VIEW DETAILS
                    if management_action == "View Details":
                        st.subheader(f"Agent Details: {selected_agent['name']}")
                        
                        col1_details, col2_details = st.columns(2)
                        
                        with col1_details:
                            st.markdown("**Basic Information:**")
                            st.json({
                                "ID": selected_agent.get("id"),
                                "Name": selected_agent.get("name"),
                                "Model": selected_agent.get("model_name"),
                                "Created": selected_agent.get("created_at"),
                                "Updated": selected_agent.get("updated_at"),
                                "Status": "Active" if selected_agent.get("is_active", True) else "Inactive"
                            })
                            
                            st.markdown("**Performance Metrics:**")
                            st.json({
                                "Total Queries": selected_agent.get("total_queries", 0),
                                "Temperature": selected_agent.get("temperature", 0.7),
                                "Max Tokens": selected_agent.get("max_tokens", 300)
                            })
                        
                        with col2_details:
                            st.markdown("**System Prompt:**")
                            st.text_area(
                                "Full System Prompt", 
                                selected_agent.get("system_prompt", ""), 
                                height=200, 
                                disabled=True,
                                key="view_system_prompt"
                            )
                            
                            st.markdown("**User Prompt Template:**")
                            st.text_area(
                                "Full User Prompt Template", 
                                selected_agent.get("user_prompt_template", ""), 
                                height=150, 
                                disabled=True,
                                key="view_user_prompt"
                            )
                    
                    # EDIT AGENT
                    elif management_action == "Edit Agent":
                        st.subheader(f"Edit Agent: {selected_agent['name']}")
                        
                        with st.form(f"edit_agent_{agent_id}"):
                            col1_edit, col2_edit = st.columns(2)
                            
                            with col1_edit:
                                new_name = st.text_input(
                                    "Agent Name", 
                                    value=selected_agent.get("name", ""),
                                    key="edit_name"
                                )
                                
                                # Model selection for edit
                                available_models = st.session_state.available_models or get_available_models_cached()
                                current_model = selected_agent.get("model_name", "")
                                
                                if current_model in available_models:
                                    model_index = available_models.index(current_model)
                                else:
                                    model_index = 0
                                
                                new_model = st.selectbox(
                                    "Model", 
                                    available_models, 
                                    index=model_index,
                                    key="edit_model"
                                )
                                
                                # Advanced settings
                                new_temperature = st.slider(
                                    "Temperature",
                                    min_value=0.0,
                                    max_value=1.0,
                                    value=selected_agent.get("temperature", 0.7),
                                    step=0.1,
                                    key="edit_temperature"
                                )
                                
                                new_max_tokens = st.number_input(
                                    "Max Tokens",
                                    min_value=100,
                                    max_value=4000,
                                    value=selected_agent.get("max_tokens", 300),
                                    step=100,
                                    key="edit_max_tokens"
                                )
                                
                                new_is_active = st.checkbox(
                                    "Active", 
                                    value=selected_agent.get("is_active", True),
                                    key="edit_active"
                                )
                            
                            with col2_edit:
                                new_system_prompt = st.text_area(
                                    "System Prompt",
                                    value=selected_agent.get("system_prompt", ""),
                                    height=200,
                                    key="edit_system_prompt"
                                )
                                
                                new_user_prompt = st.text_area(
                                    "User Prompt Template",
                                    value=selected_agent.get("user_prompt_template", ""),
                                    height=150,
                                    key="edit_user_prompt"
                                )
                            
                            # Form submission
                            submitted = st.form_submit_button("💾 Update Agent", type="primary")
                            
                            if submitted:
                                # Validation
                                if not new_name or len(new_name.strip()) < 3:
                                    st.error("Agent name must be at least 3 characters")
                                elif "{data_sample}" not in new_user_prompt:
                                    st.error("User prompt template must contain {data_sample} placeholder")
                                else:
                                    # Prepare update payload
                                    update_payload = {
                                        "name": new_name.strip(),
                                        "model_name": new_model,
                                        "system_prompt": new_system_prompt.strip(),
                                        "user_prompt_template": new_user_prompt.strip(),
                                        "temperature": new_temperature,
                                        "max_tokens": new_max_tokens,
                                        "is_active": new_is_active
                                    }
                                    
                                    try:
                                        with st.spinner("Updating agent..."):
                                            response = requests.put(
                                                f"{FASTAPI_API}/update-agent/{agent_id}",
                                                json=update_payload,
                                                timeout=30
                                            )
                                            
                                            if response.status_code == 200:
                                                st.success(f"Agent '{new_name}' updated successfully!")
                                                st.session_state.agents_data = []  # Force refresh
                                                st.rerun()
                                            else:
                                                error_detail = response.json().get("detail", response.text) if response.headers.get("content-type") == "application/json" else response.text
                                                st.error(f"Failed to update agent: {error_detail}")
                                    
                                    except Exception as e:
                                        st.error(f"Error updating agent: {str(e)}")
                    
                    # DELETE AGENT
                    elif management_action == "Delete Agent":
                        st.subheader(f"Delete Agent: {selected_agent['name']}")
                        
                        st.warning("**Permanent Action**: Agent deletion cannot be undone and will remove all associated data.")
                        
                        # Show agent info before deletion
                        with st.expander("Agent to be deleted", expanded=True):
                            st.json({
                                "ID": selected_agent.get("id"),
                                "Name": selected_agent.get("name"),
                                "Model": selected_agent.get("model_name"),
                                "Total Queries": selected_agent.get("total_queries", 0),
                                "Created": selected_agent.get("created_at")
                            })
                        
                        # Confirmation steps
                        confirm_name = st.text_input(
                            f"Type the agent name '{selected_agent['name']}' to confirm deletion:",
                            key="delete_confirm_name"
                        )
                        
                        confirm_delete = st.checkbox(
                            f"I understand this will permanently delete agent: {selected_agent['name']}",
                            key="delete_confirm_checkbox"
                        )
                        
                        if confirm_name == selected_agent['name'] and confirm_delete:
                            if st.button("Confirm Deletion", type="secondary", key="delete_confirm_button"):
                                try:
                                    with st.spinner("Deleting agent..."):
                                        response = requests.delete(f"{FASTAPI_API}/delete-agent/{agent_id}", timeout=10)
                                        
                                        if response.status_code == 200:
                                            st.success("Agent deleted successfully!")
                                            # Force refresh of agents list
                                            st.session_state.agents_data = []
                                            st.rerun()
                                        else:
                                            st.error(f"Failed to delete agent: HTTP {response.status_code}")
                                            
                                except Exception as e:
                                    st.error(f"Error deleting agent: {e}")
                        else:
                            if confirm_name != selected_agent['name'] and confirm_name:
                                st.error("Agent name doesn't match")
                            if not confirm_delete:
                                st.info("Please check the confirmation box and enter the exact agent name to enable deletion")

        else:
            # Enhanced empty state
            st.info("No agents found. Create your first agent in the 'Create New Agent' tab!")
            
            # Quick start guide
            with st.expander("Quick Start Guide"):
                st.markdown("""
                **Getting Started with Agents:**
                
                1. **Switch to 'Create New Agent'** tab above
                2. **Choose a Template**: Select from predefined templates like 'Systems Engineer' or 'Quality Control Engineer'
                3. **Configure Settings**: Adjust temperature (creativity) and max tokens (response length)
                4. **Test Configuration**: Use the test button to preview how your agent will work
                5. **Create Agent**: Click 'Create Agent' to add it to your AI toolkit
                6. **Come back here to manage**: Edit, view details, or delete agents
                """)

    # Enhanced tips and best practices (shown for both modes)
    st.markdown("---")
    with st.expander("Agent Best Practices & Tips", expanded=False):
        col1_tips, col2_tips = st.columns(2)
        
        with col1_tips:
            st.markdown("""
            **System Prompt Best Practices:**
            - Define specific expertise areas clearly
            - Include analysis frameworks and methodologies
            - Specify output format and structure
            - Add relevant standards or regulations
            - Use bullet points for clarity
            - Include examples of what to focus on
            """)
            
            st.markdown("""
            **Management Tips:**
            - Regularly review agent performance metrics
            - Update prompts based on usage patterns
            - Deactivate unused agents to keep interface clean
            - Test agents after making changes
            """)
            
        with col2_tips:
            st.markdown("""
            **User Prompt Template Tips:**
            - Always include `{data_sample}` placeholder
            - Provide clear instructions for analysis type
            - Specify desired output structure
            - Include relevant context or constraints
            - Ask for specific recommendations
            - Consider different input types (contracts, policies, etc.)
            """)
            
            st.markdown("""
            **Performance Optimization:**
            - Lower temperature (0.1-0.3) for consistent compliance checks
            - Higher temperature (0.7-0.9) for creative brainstorming
            - Adjust max tokens based on typical response needs
            - Monitor success rates and response times
            """)
        
        st.markdown("""
        **Model Selection Guidelines:**
        - **LLaMA**: Fast local processing, good for privacy-sensitive content
        
        **Temperature Settings:**
        - **0.1-0.3**: Highly consistent, factual analysis (recommended for compliance)
        - **0.4-0.7**: Balanced creativity and consistency (good for general work)
        - **0.8-1.0**: More creative responses (useful for brainstorming or alternative approaches)
        """)

    # Footer for create agent section
    st.markdown("---")
    st.warning("**Agent Disclaimer**: All created agents provide analysis for informational purposes only and do not constitute advice.")
    st.info("🔒 **Data Security**: Ensure all content processed by agents complies with your organization's data protection and confidentiality policies.")
    
    

# ----------------------------------------------------------------------
# SESSION HISTORY & ANALYTICS MODE
# ----------------------------------------------------------------------
elif chat_mode == "Session History":
    st.markdown("---")
    st.header("Agent Session History & Analytics")
    
    # Sub-mode selection
    history_mode = st.radio(
        "View:",
        ["Recent Sessions", "Analytics Dashboard", "Session Details"],
        horizontal=True
    )
    
    if history_mode == "Recent Sessions":
        st.subheader("Recent Agent Sessions")
        
        # Filters
        col1, col2, col3 = st.columns(3)
        
        with col1:
            session_limit = st.number_input("Number of sessions", min_value=10, max_value=200, value=50, step=10)
        
        with col2:
            session_type_filter = st.selectbox(
                "Filter by type:",
                ["All", "single_agent", "multi_agent_debate", "rag_analysis", "rag_debate", "compliance_check"]
            )
        
        with col3:
            if st.button("Load Sessions"):
                try:
                    # Prepare API call
                    params = {"limit": session_limit}
                    if session_type_filter != "All":
                        params["session_type"] = session_type_filter
                    
                    with st.spinner("Loading session history..."):
                        response = requests.get(f"{FASTAPI_API}/session-history", params=params, timeout=10)
                        
                        if response.status_code == 200:
                            data = response.json()
                            st.session_state.session_history = data.get("sessions", [])
                            st.success(f"Loaded {len(st.session_state.session_history)} sessions")
                        else:
                            st.error(f"Failed to load sessions: {response.status_code}")
                            
                except Exception as e:
                    st.error(f"Error: {e}")
        
        # Display sessions
        if 'session_history' in st.session_state and st.session_state.session_history:
            sessions = st.session_state.session_history
            
            # Convert to DataFrame for better display
            session_data = []
            for session in sessions:
                session_data.append({
                    "Session ID": session["session_id"][:8] + "...",
                    "Type": session["session_type"].replace("_", " ").title(),
                    "Analysis": session["analysis_type"].replace("_", " ").title(),
                    "Query Preview": session["user_query"][:100] + "..." if len(session["user_query"]) > 100 else session["user_query"],
                    "Collection": session.get("collection_name", "N/A"),
                    "Agents": session.get("agent_count", 0),
                    "Response Time": f"{session.get('total_response_time_ms', 0)}ms" if session.get('total_response_time_ms') else "N/A",
                    "Status": session["status"].title(),
                    "Created": session["created_at"][:19] if session["created_at"] else "Unknown"
                })
            
            st.dataframe(session_data, use_container_width=True, height=400)
            
            # Session details viewer
            st.subheader("View Session Details")
            session_ids = [s["session_id"] for s in sessions]
            selected_session = st.selectbox(
                "Select session to view details:",
                ["--Select Session--"] + session_ids
            )
            
            if selected_session != "--Select Session--":
                if st.button("Load Session Details"):
                    try:
                        with st.spinner("Loading session details..."):
                            response = requests.get(f"{FASTAPI_API}/session-details/{selected_session}", timeout=10)
                            
                            if response.status_code == 200:
                                details = response.json()
                                st.session_state.session_details = details
                                
                                # Display session info
                                session_info = details["session_info"]
                                
                                col1, col2 = st.columns(2)
                                with col1:
                                    st.json({
                                        "Session ID": session_info["session_id"],
                                        "Type": session_info["session_type"],
                                        "Analysis Type": session_info["analysis_type"],
                                        "Status": session_info["status"],
                                        "Agent Count": session_info["agent_count"],
                                        "Total Time": f"{session_info.get('total_response_time_ms', 0)}ms"
                                    })
                                
                                with col2:
                                    st.text_area("User Query", session_info["user_query"], height=150, disabled=True)
                                
                                # Display agent responses
                                st.subheader("🤖 Agent Responses")
                                responses = details["agent_responses"]
                                
                                for i, response in enumerate(responses, 1):
                                    sequence = response.get("sequence_order", i)
                                    agent_name = response["agent_name"]
                                    
                                    with st.expander(f"Response {sequence}: {agent_name}", expanded=(i <= 2)):
                                        col1_resp, col2_resp = st.columns([2, 1])
                                        
                                        with col1_resp:
                                            st.text_area(
                                                "Response", 
                                                response["response_text"], 
                                                height=200, 
                                                disabled=True,
                                                key=f"response_{i}"
                                            )
                                        
                                        with col2_resp:
                                            st.json({
                                                "Agent ID": response["agent_id"],
                                                "Model": response["model_used"],
                                                "Method": response["processing_method"],
                                                "Time": f"{response.get('response_time_ms', 0)}ms",
                                                "RAG Used": response.get("rag_used", False),
                                                "Docs Found": response.get("documents_found", 0)
                                            })
                            else:
                                st.error(f"Failed to load details: {response.status_code}")
                                
                    except Exception as e:
                        st.error(f"Error: {e}")
        
        else:
            st.info("Click 'Load Sessions' to view recent agent sessions")
    
    elif history_mode == "Analytics Dashboard":
        st.subheader("Analytics Dashboard")
        
        # Time period selection
        col1, col2 = st.columns([1, 3])
        
        with col1:
            days = st.selectbox("Time Period:", [1, 7, 14, 30], index=1)
        
        with col2:
            if st.button("Load Analytics"):
                try:
                    with st.spinner("Loading analytics..."):
                        response = requests.get(f"{FASTAPI_API}/session-analytics?days={days}", timeout=10)
                        
                        if response.status_code == 200:
                            analytics = response.json()
                            st.session_state.analytics = analytics
                            st.success(f"Loaded analytics for last {days} days")
                        else:
                            st.error(f"Failed to load analytics: {response.status_code}")
                            
                except Exception as e:
                    st.error(f"Error: {e}")
        
        # Display analytics
        if 'analytics' in st.session_state:
            analytics = st.session_state.analytics
            
            # Summary metrics
            col1, col2, col3, col4 = st.columns(4)
            
            session_stats = analytics["session_statistics"]
            rag_stats = analytics["rag_statistics"]
            
            with col1:
                total_sessions = sum(session_stats["by_session_type"].values())
                st.metric("Total Sessions", total_sessions)
            
            with col2:
                avg_time = session_stats["avg_response_time_ms"]
                st.metric("Avg Response Time", f"{avg_time:.0f}ms")
            
            with col3:
                st.metric("Total Responses", rag_stats["total_responses"])
            
            with col4:
                rag_rate = rag_stats["rag_usage_rate"]
                st.metric("RAG Usage Rate", f"{rag_rate:.1f}%")
            
            # Charts
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("Sessions by Type")
                session_types = session_stats["by_session_type"]
                if session_types:
                    st.bar_chart(session_types)
                else:
                    st.info("No session data available")
            
            with col2:
                st.subheader("Analysis Types")
                analysis_types = session_stats["by_analysis_type"]
                if analysis_types:
                    st.bar_chart(analysis_types)
                else:
                    st.info("No analysis data available")
            
            # Agent activity
            st.subheader("Most Active Agents")
            agent_activity = analytics["agent_activity"]
            
            if agent_activity:
                agent_data = {agent["agent_name"]: agent["response_count"] for agent in agent_activity}
                st.bar_chart(agent_data)
                
                # Detailed table
                agent_table = [
                    {
                        "Agent Name": agent["agent_name"],
                        "Agent ID": agent["agent_id"],
                        "Response Count": agent["response_count"]
                    }
                    for agent in agent_activity
                ]
                st.dataframe(agent_table, use_container_width=True)
            else:
                st.info("No agent activity data available")
        
        else:
            st.info("Click 'Load Analytics' to view performance metrics")
    
    elif history_mode == "Session Details":
        st.subheader("Search Session Details")
        
        # Manual session ID input
        session_id_input = st.text_input("Enter Session ID:")
        
        if st.button("Search Session") and session_id_input:
            try:
                with st.spinner("Searching for session..."):
                    response = requests.get(f"{FASTAPI_API}/session-details/{session_id_input}", timeout=10)
                    
                    if response.status_code == 200:
                        details = response.json()
                        
                        # Display detailed session information
                        session_info = details["session_info"]
                        
                        st.success(f"Found session: {session_info['session_type']}")
                        
                        # Session metadata
                        with st.expander("Session Information", expanded=True):
                            col1, col2 = st.columns(2)
                            
                            with col1:
                                st.json({
                                    "Session ID": session_info["session_id"],
                                    "Type": session_info["session_type"],
                                    "Analysis Type": session_info["analysis_type"],
                                    "Created": session_info["created_at"],
                                    "Completed": session_info["completed_at"],
                                    "Status": session_info["status"]
                                })
                            
                            with col2:
                                st.json({
                                    "Agent Count": session_info["agent_count"],
                                    "Total Time": f"{session_info.get('total_response_time_ms', 0)}ms",
                                    "Collection": session_info.get("collection_name", "N/A"),
                                    "Error": session_info.get("error_message", "None")
                                })
                        
                        # User query
                        st.text_area("User Query", session_info["user_query"], height=100, disabled=True)
                        
                        # Agent responses
                        responses = details["agent_responses"]
                        st.subheader(f"Agent Responses ({len(responses)})")
                        
                        for response in responses:
                            agent_name = response["agent_name"]
                            sequence = response.get("sequence_order", "N/A")
                            
                            with st.expander(f"Agent: {agent_name} (Sequence: {sequence})", expanded=False):
                                st.text_area("Response", response["response_text"], height=200, disabled=True)
                                
                                col1, col2 = st.columns(2)
                                with col1:
                                    st.json({
                                        "Processing Method": response["processing_method"],
                                        "Response Time": f"{response.get('response_time_ms', 0)}ms",
                                        "Model Used": response["model_used"]
                                    })
                                
                                with col2:
                                    st.json({
                                        "RAG Used": response.get("rag_used", False),
                                        "Documents Found": response.get("documents_found", 0),
                                        "Compliant": response.get("compliant", "N/A"),
                                        "Confidence": response.get("confidence_score", "N/A")
                                    })
                    
                    elif response.status_code == 404:
                        st.warning("Session not found")
                    else:
                        st.error(f"Error: {response.status_code}")
                        
            except Exception as e:
                st.error(f"Error: {e}")


# ----------------------------------------------------------------------
# DOCUMENT GENERATOR MODE
# ----------------------------------------------------------------------
elif chat_mode == "Document Generator":
    st.markdown("---")
    st.header("Document Generator")
    st.info("Upload templates, and generate comprehensive documentation using AI analysis.")
    
    # Document Generator sub-modes
    doc_gen_mode = st.radio(
        "Select Action:",
        ["Rule Development Agents", "Template Management", "Generate Documents", "Generated Documents"],
        horizontal=True
    )
    
    # ----------------------------------------------------------------------
    # RULE DEVELOPMENT AGENTS SUB-MODE
    # ----------------------------------------------------------------------
    if doc_gen_mode == "Rule Development Agents":
        st.subheader("Rule Development Agents")
        st.info("Create specialized agents for extracting rules, requirements, and test plans from technical documents.")
        
        # Enhanced rule development templates
        rule_agent_templates = {
            "Rule Extraction Agent": {
                "description": "Specialized agent for extracting detailed, testable rules and requirements from technical documents",
                "system_prompt": """You are a test planning expert specializing in extracting comprehensive, testable rules from technical documents.

                Your expertise includes:
                1. **Rule Identification**: Identify EVERY possible testable requirement which usually contains "shall", "must", "may", "will", "could" or "should"
                2. **Detailed Analysis**: Extract rules that are extremely detailed, explicit, and step-by-step
                3. **Measurable Criteria**: Include specific measurements, acceptable ranges, and referenced figures/tables
                4. **Test Strategy**: For ambiguous requirements, describe specific test strategies
                5. **Dependency Analysis**: Identify dependencies between rules and requirements
                6. **Conflict Detection**: Detect and resolve conflicts between requirements

                **Output Format Requirements:**
                - Use markdown headings and bolded text for organization
                - Use the provided template to structure the output
                - Include all relevant details from the analysis

                **Analysis Approach:**
                - Extract both explicit and implicit requirements
                - Consider edge cases and boundary conditions
                - Identify verification and validation methods
                - Note any missing information that would affect testing
                - Provide specific test procedures where applicable""",

                "user_prompt": """Develop a test plan based on the template provided and use the following document:

                {data_sample}

                to develop a comprehensive test plan for each requirement with a verification method, verification approach, necessary test steps.""",

                "temperature": 0.2,
                "max_tokens": 2500
            },
            
            "Test Plan Synthesis Agent": {
                "description": "Agent specialized in combining multiple rule sets into comprehensive test plans",
                "system_prompt": """You are a senior QA documentation engineer specializing in synthesizing complex test plans from multiple rule sources.

                Your expertise includes:
                1. **Test Plan Integration**: Combine multiple rule sets into coherent test plans
                2. **Cross-Reference Analysis**: Identify overlapping content and merge similar steps
                3. **Dependency Management**: Map dependencies between different sections and requirements
                4. **Conflict Resolution**: Identify and resolve conflicts between different rule sources
                5. **Test Organization**: Structure tests in logical execution order
                6. **Coverage Analysis**: Ensure comprehensive coverage of all requirements

                **Synthesis Methodology:**
                - Merge similar test steps and eliminate redundancy
                - Cross-reference overlapping content between sections
                - Organize tests by logical execution sequence
                - Identify prerequisite tests and setup requirements
                - Group related test procedures for efficiency
                - Maintain traceability to original requirements

                **Output Standards:**
                - Use content-based titles that reflect actual test scope
                - Maintain markdown formatting with clear sections
                - Provide explicit step-by-step test procedures
                - Include setup, execution, and verification steps
                - Note any special equipment or conditions required
                - Cross-reference related test procedures""",
                            
                            "user_prompt": """You are provided with detailed test rules from multiple sections. Synthesize these into a single, comprehensive test plan:

                {data_sample}

                Create a combined test plan with this structure:

                ## [Content-Based Test Plan Title]

                **Test Dependencies:**
                - List prerequisite tests and setup requirements
                - Note any equipment or environmental conditions needed

                **Conflict Resolution:**
                - Address any conflicts between different rule sources
                - Provide recommended resolution approaches

                **Integrated Test Procedures:**
                1. [Comprehensive, step-by-step test procedures]
                2. [Merge similar steps, eliminate redundancy]
                3. [Organize in logical execution order]
                4. [Include setup, execution, and verification phases]

                **Cross-References:**
                - Map relationships between different test procedures
                - Note shared requirements and common verification steps

                Focus on creating a test plan that is:
                - Logically organized and executable
                - Comprehensive in coverage
                - Efficient in execution order
                - Clear in requirements and procedures""",
            
                "temperature": 0.3,
                "max_tokens": 3000
            },
            
            "Document Section Analyzer": {
                "description": "Agent for analyzing document sections and preparing structured analysis",
                "system_prompt": """You are a technical document analysis expert specializing in structured content extraction and preparation.

                Your capabilities include:
                1. **Content Classification**: Identify types of content (requirements, procedures, specifications, etc.)
                2. **Section Analysis**: Extract key topics, themes, and technical focus areas
                3. **Structure Mapping**: Understand document hierarchy and relationships
                4. **Content Preparation**: Prepare content for further analysis by other specialized agents
                5. **Metadata Extraction**: Identify references, figures, tables, and cross-references

                **Analysis Framework:**
                - Identify the primary purpose and scope of each section
                - Extract technical specifications and requirements
                - Note procedural steps and methodologies
                - Identify measurement criteria and acceptance standards
                - Map relationships to other document sections
                - Highlight areas needing further clarification

                **Content Organization:**
                - Categorize content by type (functional, performance, interface, etc.)
                - Identify compliance requirements and standards references
                - Extract numerical values, ranges, and specifications
                - Note any conditional or situational requirements
                - Highlight critical vs. optional requirements""",
                            
                            "user_prompt": """Analyze the following document section and provide structured analysis:

                {data_sample}

                Provide analysis in this format:

                ## Section Analysis: [Content-Based Title]

                **Content Type:**
                - Identify the primary type of content (requirements, procedures, specifications, etc.)

                **Key Topics:**
                - List main topics and technical focus areas
                - Note any specialized terminology or concepts

                **Technical Specifications:**
                - Extract specific measurements, values, and criteria
                - List any referenced standards or specifications

                **Requirements Identified:**
                - Functional requirements
                - Performance requirements  
                - Interface requirements
                - Constraint requirements

                **References and Dependencies:**
                - Note any figures, tables, or cross-references mentioned
                - Identify dependencies on other sections or documents

                **Analysis Notes:**
                - Areas requiring clarification
                - Potential ambiguities or interpretation issues
                - Recommendations for further analysis

                This analysis will be used by specialized rule extraction agents.""",
                
                "temperature": 0.2,
                "max_tokens": 2000
            }
        }
        
        available_models = st.session_state.available_models or get_available_models_cached()
        
        # Manual Agent Creation Section - WITH PROPER TEMPLATE HANDLING
        st.subheader("Agent Creation")
        st.info("Create individual specialized agents with custom configurations.")
        
        # Template selection OUTSIDE the form to allow dynamic updates
        col1_pre, col2_pre = st.columns([2, 1])
        
        with col1_pre:
            selected_template = st.selectbox(
                "Choose Agent Template:",
                ["Custom"] + list(rule_agent_templates.keys()),
                key="template_selector",
                help="Select a template to auto-populate the form fields"
            )
        
        with col2_pre:
            if selected_template != "Custom":
                template_info = rule_agent_templates[selected_template]
                st.info(f"**{selected_template}**: {template_info['description']}")
        
        # Get template defaults based on selection
        if selected_template != "Custom":
            template = rule_agent_templates[selected_template]
            default_system_prompt = template["system_prompt"]
            default_user_prompt = template["user_prompt"]
            default_temp = template.get("temperature", 0.3)
            default_tokens = template.get("max_tokens", 2000)
            default_name = f"Custom {selected_template}"
        else:
            default_system_prompt = ""
            default_user_prompt = ""
            default_temp = 0.3
            default_tokens = 2000
            default_name = ""
        
        # Form with proper template integration
        with st.form("create_custom_rule_agent", clear_on_submit=False):
            col1, col2 = st.columns([1, 1])
            
            with col1:
                rule_agent_name = st.text_input(
                    "Agent Name",
                    value=default_name,
                    placeholder="e.g., Custom Standards Extractor"
                )
                
                # Model selection
                if available_models:
                    rule_agent_model = st.selectbox("Model", available_models, key="custom_model")
                else:
                    st.error("No models available")
                    rule_agent_model = None
                
                # Advanced settings in the left column
                st.subheader("Configuration")
                
                rule_temperature = st.slider(
                    "Temperature", 
                    0.0, 1.0, 
                    default_temp, 
                    0.1,
                    help="Lower = more consistent, Higher = more creative"
                )
                
                rule_max_tokens = st.number_input(
                    "Max Tokens", 
                    100, 4000, 
                    default_tokens, 
                    100,
                    help="Maximum response length"
                )
            
            with col2:
                rule_system_prompt = st.text_area(
                    "System Prompt (Define Agent's Role & Expertise)",
                    value=default_system_prompt,
                    height=300,
                    help="Define the agent's expertise and approach",
                    placeholder="You are an expert specializing in..."
                )
                
                rule_user_prompt = st.text_area(
                    "User Prompt Template",
                    value=default_user_prompt,
                    height=200,
                    help="Must include {data_sample} placeholder",
                    placeholder="Analyze the following:\n\n{data_sample}\n\nProvide analysis..."
                )
            
            # Validation indicators
            col1_val, col2_val = st.columns(2)
            
            with col1_val:
                if rule_agent_name and len(rule_agent_name.strip()) >= 3:
                    st.success("Agent name valid")
                else:
                    st.error("Agent name too short")
            
            with col2_val:
                if "{data_sample}" in rule_user_prompt:
                    st.success("User prompt has {data_sample}")
                else:
                    st.error("Missing {data_sample} placeholder")
            
            # Create button
            col1_btn, col2_btn, col3_btn = st.columns([1, 2, 1])
            
            with col2_btn:
                rule_create_submitted = st.form_submit_button(
                    "Create Custom Agent", 
                    type="primary",
                    use_container_width=True
                )
            
            if rule_create_submitted:
                # Comprehensive validation
                validation_errors = []
                
                if not rule_agent_name or len(rule_agent_name.strip()) < 3:
                    validation_errors.append("Agent name must be at least 3 characters")
                
                if not rule_system_prompt or len(rule_system_prompt.strip()) < 50:
                    validation_errors.append("System prompt must be substantial (at least 50 characters)")
                
                if "{data_sample}" not in rule_user_prompt:
                    validation_errors.append("User prompt must include {data_sample} placeholder")
                
                if not rule_agent_model:
                    validation_errors.append("Please select a valid model")
                
                if validation_errors:
                    st.error("**Validation Errors:**")
                    for error in validation_errors:
                        st.error(f"• {error}")
                else:
                    payload = {
                        "name": rule_agent_name.strip(),
                        "model_name": rule_agent_model,
                        "system_prompt": rule_system_prompt.strip(),
                        "user_prompt_template": rule_user_prompt.strip(),
                        "temperature": rule_temperature,
                        "max_tokens": rule_max_tokens
                    }
                    
                    try:
                        with st.spinner("Creating custom rule development agent..."):
                            response = requests.post(f"{FASTAPI_API}/create-agent", json=payload, timeout=30)
                            
                            if response.status_code == 200:
                                result = response.json()
                                st.success(f"Custom agent '{rule_agent_name}' created successfully!")
                                st.info(f"Agent ID: {result.get('agent_id', 'Unknown')}")
                                st.info(f"Model: {rule_agent_model}")
                                st.info(f"Temperature: {rule_temperature}")
                                
                                # Clear cached data to show new agent
                                if 'current_rule_agents' in st.session_state:
                                    del st.session_state.current_rule_agents
                                if 'agents_data' in st.session_state:
                                    st.session_state.agents_data = []
                                
                            else:
                                error_detail = response.json().get("detail", response.text) if response.headers.get("content-type") == "application/json" else response.text
                                
                                if "already exists" in error_detail.lower():
                                    st.error(f"Agent name '{rule_agent_name}' already exists. Please choose a different name.")
                                else:
                                    st.error(f"Failed to create agent: {error_detail}")
                    except Exception as e:
                        st.error(f"Error: {str(e)}")
        
        st.markdown("---")
        
        # Agent Management Section
        st.subheader("Manage Rule Development Agents")
        
        col1_mgmt, col2_mgmt = st.columns([1, 1])
        
        with col1_mgmt:
            if st.button("Load All Agents", key="load_all_agents"):
                try:
                    with st.spinner("Loading all agents..."):
                        agents_response = requests.get(f"{FASTAPI_API}/get-agents", timeout=10)
                        if agents_response.status_code == 200:
                            all_agents = agents_response.json().get("agents", [])
                            st.session_state.current_rule_agents = all_agents
                            st.success(f"Loaded {len(all_agents)} total agents")
                        else:
                            st.error("Failed to load agents")
                except Exception as e:
                    st.error(f"Error: {e}")
        
        with col2_mgmt:
            if 'current_rule_agents' in st.session_state:
                total_agents = len(st.session_state.current_rule_agents)
                rule_related = sum(1 for agent in st.session_state.current_rule_agents 
                                if any(keyword in agent['name'].lower() 
                                    for keyword in ['rule', 'test', 'document', 'analysis', 'extract', 'auto']))
                st.metric("Total Agents", total_agents, delta=f"{rule_related} rule-related")
        
        # Display current agents
        if 'current_rule_agents' in st.session_state and st.session_state.current_rule_agents:
            agents = st.session_state.current_rule_agents
            
            # Filter for rule development related agents
            rule_agents = [agent for agent in agents 
            if any(keyword in agent['name'].lower() for keyword in ['rule', 'test', 'document', 'analysis', 'extract', 'synthesis', 'auto'])]
            
            if rule_agents:
                st.subheader(f"Rule Development Agents ({len(rule_agents)} found)")
                
                # Create enhanced display table
                agent_display_data = []
                for agent in rule_agents:
                    agent_display_data.append({
                        "Name": agent.get("name", "Unknown"),
                        "Model": agent.get("model_name", "Unknown"),
                        "Temperature": f"{agent.get('temperature', 0.7):.1f}",
                        "Max Tokens": agent.get("max_tokens", 1000),
                        "Queries": agent.get("total_queries", 0),
                        "Status": "Active" if agent.get("is_active", True) else "Inactive",
                        "Created": agent.get("created_at", "Unknown")[:10] if agent.get("created_at") else "Unknown"
                    })
                
                st.dataframe(agent_display_data, use_container_width=True, height=300)
                
                # Quick Test Section
                st.subheader("Test Agent")
                
                if rule_agents:
                    test_agent_choice = st.selectbox(
                        "Select agent to test:",
                        [f"{agent['name']} (ID: {agent['id']})" for agent in rule_agents],
                        key="test_agent_select"
                    )
                    
                    # Sample test content
                    test_content = st.text_area(
                        "Test Content:",
                        value="""4.2.3 Signal Processing Requirements
                        The system SHALL process incoming RF signals according to the following specifications:
                        4.2.3.1 Frequency Range: The system SHALL operate within the frequency range of 30 MHz to 3 GHz with a tolerance of ±0.1%.
                        4.2.3.2 Signal Sensitivity: The minimum detectable signal level SHALL be -110 dBm or better across the entire frequency range.
                        4.2.3.3 Processing Time: Signal processing SHALL be completed within 50 milliseconds from signal acquisition to output generation.""",
                        height=200,
                        help="Edit this content to test your agent"
                    )
                    
                    if st.button("Run Test", type="secondary"):
                        if test_agent_choice and test_content:
                            # Extract agent ID from selection
                            agent_id = int(test_agent_choice.split("ID: ")[1].rstrip(")"))
                            agent_name = test_agent_choice.split(" (ID:")[0]
                            
                            with st.spinner(f"Testing {agent_name}..."):
                                try:
                                    payload = {
                                        "data_sample": test_content,
                                        "agent_ids": [agent_id]
                                    }
                                    
                                    response = requests.post(
                                        f"{FASTAPI_API}/compliance-check",
                                        json=payload,
                                        timeout=60
                                    )
                                    
                                    if response.status_code == 200:
                                        result = response.json()
                                        st.success("Test completed successfully!")
                                        
                                        # Show test results
                                        details = result.get("details", {})
                                        for idx, analysis in details.items():
                                            st.subheader(f"Results from {analysis.get('agent_name', 'Unknown Agent')}")
                                            st.markdown(analysis.get("reason", "No analysis generated"))
                                            
                                            if "response_time_ms" in result:
                                                st.caption(f"Response time: {result['response_time_ms']}ms")
                                    else:
                                        st.error(f"Test failed: HTTP {response.status_code}")
                                        error_detail = response.json().get("detail", response.text) if response.headers.get("content-type") == "application/json" else response.text
                                        st.error(f"Error details: {error_detail}")
                                        
                                except Exception as e:
                                    st.error(f"Test error: {str(e)}")
            
            else:
                st.info("No rule development agents found. Create some using the options above!")
        
        else:
            st.info("Click 'Load All Agents' to see existing agents and their status.")
        
        # Footer
        st.markdown("---")
        st.info("**Next Steps**: After creating agents, use them in 'Generate Documents' for document analysis!")
    
    
    # ----------------------------------------------------------------------
    # TEMPLATE MANAGEMENT SUB-MODE  
    # ----------------------------------------------------------------------
    
    if doc_gen_mode == "Template Management":
        st.header("Document Template Management")
        st.info("Upload and manage document templates for automated rule generation and test plan creation.")
        
        collections = get_chromadb_collections()

        render_upload_component(
            available_collections= collections,
            load_collections_func= get_chromadb_collections,
            create_collection_func= create_collection,
            upload_endpoint=f"{CHROMADB_API}/documents/upload-and-process",
            job_status_endpoint=f"{CHROMADB_API}/jobs/{{job_id}}"
        )
        
    # ----------------------------------------------------------------------
    # GENERATE DOCUMENTS SUB-MODE
    # ----------------------------------------------------------------------
    elif doc_gen_mode == "Generate Documents":
        st.subheader("Generate Test Plans")

        # ——————————————————————————
        # 1) Pick agents
        # ——————————————————————————
        agents = st.session_state.get("available_rule_agents") or requests.get(f"{FASTAPI_API}/get-agents").json()["agents"]
        # rule_agents = [a for a in agents if "rule" in a["name"].lower()]
        agent_map = {f"{a['name']} ({a['model_name']})": a["id"] for a in agents}
        selected_agents = st.multiselect("Select Agents", list(agent_map.keys()), key="gen_agents")
        if not selected_agents:
            st.info("Choose at least one agent to proceed"); st.stop()

        # ——————————————————————————
        # 2a) Pick TEMPLATE collection & load template docs
        # ——————————————————————————
        template_collection = st.selectbox(
            "Template Collection (the one you uploaded as templates)",
            st.session_state.collections,
            key="gen_template_coll",
        )
        if st.button("Load Template Library", key="gen_load_templates"):
            st.session_state.template_docs = get_all_documents_in_collection(template_collection)

        template_docs = st.session_state.get("template_docs", [])
        template_map = {d["document_name"]: d["document_id"] for d in template_docs}
        selected_templates = st.multiselect(
            "Select Template(s)", list(template_map.keys()), key="gen_templates"
        )
        template_doc_ids = [template_map[name] for name in selected_templates]

        # ——————————————————————————
        # 2b) Pick SOURCE collection & load source docs
        # ——————————————————————————
        source_collection = st.selectbox(
            "Source Collection (your requirements/standards)",
            st.session_state.collections,
            key="gen_source_coll",
        )
        if st.button("Load Source Documents", key="gen_load_sources"):
            st.session_state.source_docs = get_all_documents_in_collection(source_collection)

        source_docs = st.session_state.get("source_docs", [])
        source_map = {d["document_name"]: d["document_id"] for d in source_docs}
        selected_sources = st.multiselect(
            "Select Source Document(s)", list(source_map.keys()), key="gen_sources"
        )
        source_doc_ids = [source_map[name] for name in selected_sources]

        # ——————————————————————————
        # 3) Pick agents 
        # ——————————————————————————
        agent_ids = [agent_map[label] for label in selected_agents]

        # ——————————————————————————
        # 4) Let user name the output file
        # ——————————————————————————
        out_name = st.text_input(
            "Output file name (no extension):",
            value="Generated_Analysis",
            key="gen_filename"
        ).strip()

        # ——————————————————————————
        # 5) Generate analyses
        # ——————————————————————————
        if st.button("Generate Documents", type="primary"):
            if not template_doc_ids or not source_doc_ids:
                st.error("You must select at least one templated and one source doc.")
            else:
                payload = {
                    "template_collection": template_collection,
                    "template_doc_ids":    template_doc_ids,
                    "source_collections":  [source_collection],
                    "source_doc_ids":      source_doc_ids,
                    "agent_ids":           agent_ids,
                    "use_rag":             True,
                    "top_k":               5,
                }
                st.write("about to call /generate_documents on", FASTAPI_API)
                st.write("Payload:", payload)
                with st.spinner("Calling Document Generator…"):
                    try:
                        resp = requests.post(
                            f"{FASTAPI_API}/generate_documents",
                            json=payload,
                            timeout=300
                        )
                        # now resp is guaranteed to exist
                        if not resp.ok:
                            st.error(f"Error {resp.status_code}: {resp.text}")
                        else:
                            docs = resp.json().get("documents", [])
                            st.success(f"Generated {len(docs)} documents")
                            for d in docs:
                                blob = base64.b64decode(d["docx_b64"])
                                st.download_button(
                                    label=f"{d['title']}.docx",
                                    data=blob,
                                    file_name=f"{d['title']}.docx",
                                    mime=(
                                        "application/"
                                        "vnd.openxmlformats-"
                                        "officedocument."
                                        "wordprocessingml.document"
                                    )
                                )
                    except Exception as e:
                            st.error("Request exception: " + str(e))

        # ——————————————————————————
        # 6) Offer download once we have results
        # ——————————————————————————
        if st.session_state.get("gen_results"):
            buffer = build_docx_bytes([
                {
                    "document_title": r["title"],
                    "analysis_content": r["content"],
                    "source_document": "",
                    "agent_name": ""
                }
                for r in st.session_state.gen_results
            ])
            st.download_button(
                "Download Combined DOCX",
                data=buffer.getvalue(),
                file_name=f"{out_name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )


    # ----------------------------------------------------------------------
    # Sub GENERATED DOCUMENTS SUBMODE
    # ----------------------------------------------------------------------

    elif doc_gen_mode == "Generated Documents":
        st.subheader("Generated Test Plans")
        st.info("View, manage, and export your generated rule analysis documents.")
        
        # Check if there are generated results in session state
        if 'generated_results' in st.session_state and st.session_state.generated_results:
            results = st.session_state.generated_results
            output_collection = st.session_state.get('output_collection', 'generated-documents')
            
            st.success(f"Found {len(results)} generated documents from your last generation session")
            
            # Results Overview
            st.subheader("Generation Overview")
            
            col1_overview, col2_overview, col3_overview, col4_overview = st.columns(4)
            
            with col1_overview:
                st.metric("Total Documents", len(results))
            
            with col2_overview:
                unique_agents = len(set(r['agent_name'] for r in results))
                st.metric("Agents Used", unique_agents)
            
            with col3_overview:
                unique_sources = len(set(r['source_document'] for r in results))
                st.metric("Source Documents", unique_sources)
            
            with col4_overview:
                total_chars = sum(r['content_length'] for r in results)
                st.metric("Total Content", f"{total_chars:,} chars")
            
            # Detailed Results Table
            st.subheader("Generated Documents")
            
            # Create results table
            results_table = []
            for i, result in enumerate(results):
                results_table.append({
                    "Index": i + 1,
                    "Document Title": result['document_title'],
                    "Source Document": result['source_document'],
                    "Agent": result['agent_name'],
                    "Content Length": f"{result['content_length']:,} chars",
                    "Generated": result['generation_timestamp'][:19],
                    "Processing Time": f"{result.get('processing_time_ms', 0)}ms"
                })
            
            st.dataframe(results_table, use_container_width=True, height=400)
            
            # Document Actions
            st.markdown("---")
            st.subheader("Document Actions")
            
            # Document selector for individual actions
            doc_titles = [r['document_title'] for r in results]
            selected_doc_title = st.selectbox(
                "Select document for actions:",
                ["--Select Document--"] + doc_titles,
                key="selected_generated_doc"
            )
            
            if selected_doc_title != "--Select Document--":
                # Find the selected result
                selected_result = next((r for r in results if r['document_title'] == selected_doc_title), None)
                
                if selected_result:
                    col1_action, col2_action, col3_action = st.columns(3)
                    
                    with col1_action:
                        if st.button("Preview Document", key="preview_generated_doc"):
                            st.subheader(f"Preview: {selected_result['document_title']}")
                            
                            # Document metadata
                            with st.expander("Document Information", expanded=False):
                                st.json({
                                    "Source Document": selected_result['source_document'],
                                    "Agent": selected_result['agent_name'],
                                    "Generated": selected_result['generation_timestamp'],
                                    "Content Length": f"{selected_result['content_length']:,} characters",
                                    "Processing Time": f"{selected_result.get('processing_time_ms', 0)}ms"
                                })
                            
                            # Content preview
                            content = selected_result['analysis_content']
                            st.text_area(
                                "Generated Content:",
                                content,
                                height=400,
                                disabled=True,
                                key="preview_content_area"
                            )
                    
                    with col2_action:
                        if st.button("Download as DOCX", key="download_single_docx"):
                            try:
                                # Create DOCX document
                                doc = Document()
                                doc.add_heading(selected_result['document_title'], 0)
                                
                                # Add metadata section
                                doc.add_heading('Document Information', 1)
                                doc.add_paragraph(f"Source Document: {selected_result['source_document']}")
                                doc.add_paragraph(f"Generated by Agent: {selected_result['agent_name']}")
                                doc.add_paragraph(f"Generated on: {selected_result['generation_timestamp']}")
                                doc.add_paragraph(f"Content Length: {selected_result['content_length']:,} characters")
                                
                                # Add main content
                                doc.add_heading('Analysis Content', 1)
                                
                                content = selected_result['analysis_content']
                                markdown_to_docx(content, doc)
                                
                                # Save to temporary file
                                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                                doc.save(temp_file.name)
                                
                                # Create download
                                with open(temp_file.name, "rb") as file:
                                    st.download_button(
                                        label="Download DOCX",
                                        data=file.read(),
                                        file_name=f"{selected_result['document_title'].replace(' ', '_')}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                        key="download_single_docx_btn"
                                    )
                                
                                # Cleanup
                                os.unlink(temp_file.name)
                                st.success("Download ready!")
                                
                            except Exception as e:
                                st.error(f"Download error: {e}")
                    
                    with col3_action:
                        if st.button("Save to Collection", key="save_to_collection"):
                            try:
                                # Create a temporary file with the content
                                temp_file = tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".txt", encoding='utf-8')
                                temp_file.write(selected_result['analysis_content'])
                                temp_file.close()
                                
                                # Read the file back as bytes for upload
                                with open(temp_file.name, 'rb') as file:
                                    file_data = file.read()
                                
                                # Create a file-like object for the upload
                                from io import BytesIO
                                file_obj = BytesIO(file_data)
                                file_obj.name = f"{selected_result['document_title']}.txt"
                                
                                # Store in ChromaDB
                                result = store_files_in_chromadb(
                                    [file_obj], 
                                    output_collection,
                                    model_name="basic",
                                    chunk_size=2000,
                                    chunk_overlap=200,
                                    store_images=False
                                )
                                
                                st.success(f"Document saved to collection '{output_collection}'!")
                                st.json(result)
                                
                                # Cleanup
                                os.unlink(temp_file.name)
                                
                            except Exception as e:
                                st.error(f"Save error: {e}")
            
            st.markdown("---")
            
            # Bulk Actions
            st.subheader("Bulk Actions")
            
            col1_bulk, col2_bulk, col3_bulk = st.columns(3)
            
            with col1_bulk:
                if st.button("Download All as Combined DOCX", key="download_all_combined"):
                    try:
                        from docx import Document
                        import tempfile
                        import os
                        
                        # Create combined document
                        doc = Document()
                        doc.add_heading('Generated Rule Analysis Documents', 0)
                        
                        # Add generation summary
                        doc.add_heading('Generation Summary', 1)
                        doc.add_paragraph(f"Total Documents: {len(results)}")
                        doc.add_paragraph(f"Agents Used: {unique_agents}")
                        doc.add_paragraph(f"Source Documents: {unique_sources}")
                        doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                        doc.add_page_break()
                        
                        # Add each document
                        for i, result in enumerate(results, 1):
                            doc.add_heading(f"{i}. {result['document_title']}", 1)
                            
                            # Add metadata
                            doc.add_paragraph(f"Source: {result['source_document']}")
                            doc.add_paragraph(f"Agent: {result['agent_name']}")
                            doc.add_paragraph(f"Generated: {result['generation_timestamp'][:19]}")
                            doc.add_paragraph("")  # Empty line
                            
                            # Add content
                            content = result['analysis_content']
                            markdown_to_docx(content, doc)
                            
                            if i < len(results):  # Add page break except for last document
                                doc.add_page_break()
                        
                        # Save and provide download
                        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                        doc.save(temp_file.name)
                        
                        with open(temp_file.name, "rb") as file:
                            st.download_button(
                                label="Download Combined DOCX",
                                data=file.read(),
                                file_name=f"Generated_Analysis_Combined_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="download_combined_docx_btn"
                            )
                        
                        os.unlink(temp_file.name)
                        st.success("Combined download ready!")
                        
                    except Exception as e:
                        st.error(f"Combined download error: {e}")
            
            with col2_bulk:
                if st.button("Save All to Collection", key="save_all_to_collection"):
                    try:
                        import tempfile
                        import os
                        from io import BytesIO
                        
                        saved_count = 0
                        failed_count = 0
                        
                        with st.spinner(f"Saving {len(results)} documents to collection '{output_collection}'..."):
                            progress_bar = st.progress(0)
                            
                            for i, result in enumerate(results):
                                try:
                                    # Create temporary file
                                    temp_file = tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".txt", encoding='utf-8')
                                    temp_file.write(result['analysis_content'])
                                    temp_file.close()
                                    
                                    # Read back as bytes
                                    with open(temp_file.name, 'rb') as file:
                                        file_data = file.read()
                                    
                                    # Create file object
                                    file_obj = BytesIO(file_data)
                                    file_obj.name = f"{result['document_title']}.txt"
                                    
                                    # Store in ChromaDB
                                    store_files_in_chromadb(
                                        [file_obj], 
                                        output_collection,
                                        model_name="basic",
                                        chunk_size=2000,
                                        chunk_overlap=200,
                                        store_images=False
                                    )
                                    
                                    saved_count += 1
                                    os.unlink(temp_file.name)
                                    
                                except Exception as e:
                                    failed_count += 1
                                    st.warning(f"Failed to save '{result['document_title']}': {e}")
                                
                                progress_bar.progress((i + 1) / len(results))
                        
                        if saved_count > 0:
                            st.success(f"Successfully saved {saved_count} documents to collection '{output_collection}'!")
                        if failed_count > 0:
                            st.warning(f"{failed_count} documents failed to save.")
                        
                    except Exception as e:
                        st.error(f"Bulk save error: {e}")
            
            with col3_bulk:
                if st.button("Clear Generated Results", key="clear_results"):
                    if st.button("Confirm Clear", key="confirm_clear"):
                        del st.session_state.generated_results
                        if 'output_collection' in st.session_state:
                            del st.session_state.output_collection
                        st.success("Generated results cleared!")
                        st.rerun()
                    else:
                        st.warning("Click 'Confirm Clear' to permanently remove results from session")
        
        else:
            # No generated results found
            st.info("No generated documents found in the current session.")
            
            # Check for existing collections with generated content
            st.subheader("Browse Existing Generated Documents")
            
            current_collections = st.session_state.collections or []
            
            if current_collections:
                # Filter for likely generated document collections
                generated_collections = [col for col in current_collections 
                                    if any(keyword in col.lower() 
                                            for keyword in ['generated', 'analysis', 'output', 'result', 'rule'])]
                
                if generated_collections:
                    st.write("**Collections that might contain generated documents:**")
                    
                    browse_collection = st.selectbox(
                        "Select collection to browse:",
                        generated_collections,
                        key="browse_generated_collection"
                    )
                    
                    if st.button("Load Generated Documents", key="load_generated_docs"):
                        try:
                            with st.spinner("Loading documents from collection..."):
                                documents = get_all_documents_in_collection(browse_collection)
                                
                                if documents:
                                    st.success(f"Found {len(documents)} documents in '{browse_collection}'")
                                    
                                    # Display documents
                                    doc_overview = []
                                    for doc in documents:
                                        doc_overview.append({
                                            "Document Name": doc["document_name"],
                                            "File Type": doc["file_type"].upper(),
                                            "Chunks": doc["total_chunks"],
                                            "Uploaded": doc["processing_timestamp"][:10] if doc["processing_timestamp"] else "Unknown",
                                            "Document ID": doc["document_id"][:12] + "..."
                                        })
                                    
                                    st.dataframe(doc_overview, use_container_width=True)
                                    
                                    # Document selector for actions
                                    doc_choices = {f"{doc['document_name']} ({doc['document_id'][:8]}...)": doc['document_id'] for doc in documents}
                                    selected_existing = st.selectbox(
                                        "Select document to view/download:",
                                        ["--Select Document--"] + list(doc_choices.keys()),
                                        key="selected_existing_doc"
                                    )
                                    
                                    if selected_existing != "--Select Document--":
                                        doc_id = doc_choices[selected_existing]
                                        
                                        col1_existing, col2_existing = st.columns(2)
                                        
                                        with col1_existing:
                                            if st.button("Preview Document", key="preview_existing_doc"):
                                                try:
                                                    with st.spinner("Loading document..."):
                                                        result = reconstruct_document_with_timeout(doc_id, browse_collection, timeout=120)
                                                        
                                                        st.subheader(f"📄 {result['document_name']}")
                                                        
                                                        content = result['reconstructed_content']
                                                        st.text_area(
                                                            f"Content ({len(content):,} characters):",
                                                            content,
                                                            height=400,
                                                            disabled=True,
                                                            key="existing_content_preview"
                                                        )
                                                        
                                                except Exception as e:
                                                    st.error(f"Preview error: {e}")
                                        
                                        with col2_existing:
                                            if st.button("Download as DOCX", key="download_existing_docx"):
                                                try:
                                                    with st.spinner("Preparing download..."):
                                                        result = reconstruct_document_with_timeout(doc_id, browse_collection, timeout=120)
                                                        
                                                        # Create DOCX
                                                        docx_path = export_to_docx(result)
                                                        
                                                        with open(docx_path, "rb") as file:
                                                            st.download_button(
                                                                label="Download DOCX",
                                                                data=file.read(),
                                                                file_name=f"{result['document_name']}.docx",
                                                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                                                key="download_existing_docx_btn"
                                                            )
                                                        
                                                        os.unlink(docx_path)
                                                        st.success("Download ready!")
                                                        
                                                except Exception as e:
                                                    st.error(f"Download error: {e}")
                                else:
                                    st.info(f"No documents found in collection '{browse_collection}'")
                                    
                        except Exception as e:
                            st.error(f"Error loading collection: {e}")
                
                else:
                    st.info("No collections found that appear to contain generated documents.")
            
            else:
                st.warning("No collections available.")
            

# Footer
st.markdown("---")
st.caption("This application processes documents and provide GenAI capabilitites. Ensure all data is handled according to your organization's data protection policies.")
